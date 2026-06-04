# -*- coding: utf-8 -*-
"""演示 KL 标准文档构建、分类和检查。"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from docxspec import build_standard_docx, check_word_standard, classify_docx_body


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "output"
TEMPLATE = ROOT / "template.docx"


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    body_docx = OUTPUT / "demo10_kl_body.docx"
    output_docx = OUTPUT / "demo10_kl_standard_output.docx"
    report_json = OUTPUT / "demo10_kl_report.json"

    doc = Document()
    doc.add_paragraph("积分元件测试大纲")
    doc.add_paragraph("1 需求背景")
    doc.add_paragraph("这是待规范化的正文内容。")
    doc.add_paragraph("1.1 测试范围")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "名称"
    table.cell(1, 1).text = "积分元件"
    doc.save(body_docx)

    build_report = build_standard_docx(body_docx, output_docx, TEMPLATE)
    classify_report = classify_docx_body(output_docx)
    check_report = check_word_standard(output_docx)

    report_json.write_text(
        json.dumps(
            {
                "build": build_report,
                "classification_summary": classify_report["summary"],
                "check": check_report,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    print(f"已生成标准文档: {output_docx}")
    print(f"已生成检查报告: {report_json}")


if __name__ == "__main__":
    main()
