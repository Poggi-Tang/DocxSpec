from pathlib import Path
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from docxspec import WordAPI


DEMO_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = DEMO_DIR / "output"
TEMPLATE_PATH = OUTPUT_DIR / "demo9_block_template_source.docx"
OUTPUT_PATH = OUTPUT_DIR / "demo9_block_template_output.docx"


def add_seq_field(paragraph, field_code: str, result: str = "1") -> None:
    """向段落中写入一个 Word SEQ 域，用于模拟真实表题注。"""
    begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin._r.append(fld_begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    instr_run._r.append(instr)

    sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep._r.append(fld_sep)

    paragraph.add_run(result)

    end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end._r.append(fld_end)


def create_demo_template() -> None:
    """创建一个最小可运行模板，包含来源块标记和容器占位符。"""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("BlockTemplate Demo", level=1)
    doc.add_paragraph("{{p result}}")

    doc.add_paragraph("{{SOURCE_TABLE_BLOCK}}")
    caption = doc.add_paragraph("表 ")
    add_seq_field(caption, r"SEQ 表 \* ARABIC", "1")
    caption.add_run(" {{TABLE_TITLE}}")

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "测试项标识"
    table.cell(0, 1).text = "{{ITEM_ID}}"
    table.cell(1, 0).text = "测试项名称"
    table.cell(1, 1).text = "{{ITEM_NAME}}"
    table.cell(2, 0).text = "测试目的"
    table.cell(2, 1).text = "{{PURPOSE}}"

    doc.save(TEMPLATE_PATH)


def main() -> Path:
    """演示从 Word 中抽取“题注+表格”块，再放入容器重复渲染。"""
    create_demo_template()

    api = WordAPI(str(TEMPLATE_PATH))
    block_template = api.extract_table_block("{{SOURCE_TABLE_BLOCK}}", remove_block=True)

    result = api.new_container()
    rows = [
        ("FT_Demo_UI_001", "默认图标显示功能", "验证元件默认图标显示是否正确"),
        ("FT_Demo_Param_001", "参数配置功能", "验证参数配置是否能正确生效"),
    ]
    for item_id, item_name, purpose in rows:
        block = block_template.clone().replace_text(
            {
                "{{TABLE_TITLE}}": item_name,
                "{{ITEM_ID}}": item_id,
                "{{ITEM_NAME}}": item_name,
                "{{PURPOSE}}": purpose,
            }
        )
        result.add_block(block)

    api.render({"result": result.subdoc}, str(OUTPUT_PATH))
    return OUTPUT_PATH


if __name__ == "__main__":
    print(f"Generated: {main()}")
