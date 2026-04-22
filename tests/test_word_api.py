# -*- coding: utf-8 -*-
"""
    docxspec WordAPI 完整功能测试套件
"""

import os
from pathlib import Path

import pytest
from docx import Document

from docxspec import WordAPI, make_rich_text, make_table_style
from docxspec.word_styles import (
    BODY_STYLE,
    CAPTION_STYLE,
    IMAGE_STYLE,
    TABLE_BODY_STYLE,
    TABLE_HEADER_STYLE,
    make_cell_style,
)


@pytest.fixture
def template_path():
    """获取测试模板文件路径。"""
    return str(Path(__file__).parent / "templates" / "template.docx")


@pytest.fixture
def test_image():
    """获取测试图片文件路径。"""
    return str(Path(__file__).parent / "templates" / "docxspec.png")


@pytest.fixture
def output_dir():
    """创建测试输出目录。"""
    path = Path(__file__).parent / "output"
    path.mkdir(parents=True, exist_ok=True)
    return path


def build_min_context(api: WordAPI):
    return {
        "text_tag": make_rich_text("测试", BODY_STYLE),
        "image_tag": api.new_container().subdoc,
        "table_tag": api.new_container().subdoc,
        "result": api.new_container().subdoc,
    }


class TestWordAPIAllInOne:
    """统一测试类：保留尽量完整的测试项，不拆多个测试类。"""

    # =========================
    # 1. 初始化与基础能力
    # =========================
    def test_init_with_valid_template(self, template_path):
        api = WordAPI(template_path)
        assert api.template_path == template_path

    def test_init_with_invalid_template(self):
        with pytest.raises(FileNotFoundError):
            WordAPI("nonexistent_template.docx")

    # =========================
    # 2. DocContainer 链式能力
    # =========================
    def test_container_add_title(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        result = container.add_title("测试标题")
        assert result is container

    def test_container_add_heading(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        for level in [1, 2, 3]:
            result = container.add_heading(f"{level}级标题", level=level)
            assert result is container

    def test_container_add_paragraph(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        result = container.add_paragraph("测试段落")
        assert result is container

    def test_container_add_image(self, template_path, test_image):
        api = WordAPI(template_path)
        container = api.new_container()
        result = container.add_image(test_image, width_cm=6.0)
        assert result is container

    def test_container_add_table(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        data = [["表头1", "表头2"], ["数据1", "数据2"]]
        result = container.add_table(data)
        assert result is container

    def test_container_add_table_by_config(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        config = {
            "data": [["序号", "内容"], ["1", "测试"]],
            "style": {
                "header": TABLE_HEADER_STYLE,
                "body": TABLE_BODY_STYLE,
            },
        }
        result = container.add_table_by_config(config)
        assert result is container

    def test_container_add_page_break(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        result = container.add_page_break()
        assert result is container

    def test_container_add_field_paragraph(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        parts = [
            {"type": "text", "value": "第 "},
            {"type": "field", "code": "PAGE"},
            {"type": "text", "value": " 页"},
        ]
        result = container.add_field_paragraph(parts)
        assert result is not None

    def test_container_add_figure_caption_auto(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        result = container.add_figure_caption_auto("示例图")
        assert result is not None

    def test_container_add_table_caption_auto(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        result = container.add_table_caption_auto("示例表")
        assert result is not None

    def test_container_chain_calling(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        result = (
            container
            .add_title("主标题")
            .add_heading("一级标题", level=1)
            .add_paragraph("正文内容")
            .add_page_break()
        )
        assert result is container

    # =========================
    # 3. 段落能力
    # =========================
    def test_add_empty_paragraph(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        paragraph = api.add_empty_paragraph(container.subdoc, BODY_STYLE)
        assert paragraph is not None

    def test_add_text_run(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        paragraph = container.subdoc.add_paragraph()
        run = api.add_text_run(paragraph, "测试文本", BODY_STYLE)
        assert run is not None

    def test_add_paragraph_with_text(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        paragraph = api.add_paragraph(container.subdoc, "完整段落", BODY_STYLE)
        assert paragraph is not None

    def test_add_paragraph_with_none_text(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        paragraph = api.add_paragraph(container.subdoc, None, BODY_STYLE)
        assert paragraph is not None

    # =========================
    # 4. 图片能力
    # =========================
    def test_add_image_with_width(self, template_path, test_image):
        api = WordAPI(template_path)
        container = api.new_container()
        paragraph = api.add_image_block(
            container.subdoc,
            test_image,
            width_cm=6.0,
            align="center",
        )
        assert paragraph is not None

    def test_add_image_with_height(self, template_path, test_image):
        api = WordAPI(template_path)
        container = api.new_container()
        paragraph = api.add_image_block(
            container.subdoc,
            test_image,
            height_cm=4.0,
            align="left",
        )
        assert paragraph is not None

    def test_add_image_with_both_dimensions(self, template_path, test_image):
        api = WordAPI(template_path)
        container = api.new_container()
        paragraph = api.add_image_block(
            container.subdoc,
            test_image,
            width_cm=6.0,
            height_cm=4.0,
            align="right",
        )
        assert paragraph is not None

    def test_add_image_without_dimensions(self, template_path, test_image):
        api = WordAPI(template_path)
        container = api.new_container()
        paragraph = api.add_image_block(
            container.subdoc,
            test_image,
            style=IMAGE_STYLE,
        )
        assert paragraph is not None

    def test_add_image_not_found(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        with pytest.raises(FileNotFoundError):
            api.add_image_block(container.subdoc, "nonexistent.png")

    # =========================
    # 5. 表格能力
    # =========================
    def test_insert_basic_table(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        data = [
            ["姓名", "年龄", "城市"],
            ["张三", "25", "北京"],
            ["李四", "30", "上海"],
        ]
        table = api.insert_table(container.subdoc, data)
        assert table is not None
        assert len(table.rows) == 3

    def test_insert_table_with_custom_styles(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        data = [["A", "B"], ["1", "2"]]
        header_style = make_cell_style(bg_color="FF0000", bold=True)
        body_style = make_cell_style(font_size=10)
        table_style = make_table_style(border_color="0000FF")
        table = api.insert_table(
            container.subdoc,
            data,
            header_style=header_style,
            body_style=body_style,
            table_style=table_style,
        )
        assert table is not None

    def test_insert_table_empty_data(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        with pytest.raises(ValueError):
            api.insert_table(container.subdoc, [])

    def test_insert_table_irregular_rows(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        data = [["A", "B", "C"], ["1", "2"], ["X"]]
        table = api.insert_table(container.subdoc, data)
        assert table is not None

    def test_insert_table_with_col_widths(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        data = [["列1", "列2", "列3"], ["a", "b", "c"]]
        table_style = make_table_style(col_widths_cm=[3.0, 4.0, 5.0])
        table = api.insert_table(container.subdoc, data, table_style=table_style)
        assert table is not None

    def test_insert_table_with_row_heights(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        data = [["行1"], ["行2"]]
        table_style = make_table_style(row_heights_cm=[1.0, 1.5])
        table = api.insert_table(container.subdoc, data, table_style=table_style)
        assert table is not None

    def test_insert_table_with_image_cell(self, template_path, test_image):
        api = WordAPI(template_path)
        container = api.new_container()
        data = [["图片", "说明"], [test_image, "表格中的图片"]]
        table = api.insert_table(container.subdoc, data)
        assert table is not None
        cell_xml = table.cell(1, 0)._tc.xml
        assert "graphicData" in cell_xml or "pic:pic" in cell_xml

    # =========================
    # 6. 配置表格能力
    # =========================
    def test_insert_table_by_config_full(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        config = {
            "data": [["序号", "名称"], ["1", "项目A"]],
            "row_heights_cm": [0.8, 0.6],
            "col_widths_cm": [2.0, 6.0],
            "style": {
                "header": TABLE_HEADER_STYLE,
                "body": TABLE_BODY_STYLE,
                "table": make_table_style(),
            },
        }
        table = api.insert_table_by_config(container.subdoc, config)
        assert table is not None

    def test_insert_table_by_config_minimal(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        config = {"data": [["简单表格"]]}
        table = api.insert_table_by_config(container.subdoc, config)
        assert table is not None

    def test_insert_table_by_config_empty_data(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        config = {}
        table = api.insert_table_by_config(container.subdoc, config)
        assert table is not None

    # =========================
    # 7. 域与题注
    # =========================
    def test_add_field_run(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        paragraph = container.subdoc.add_paragraph()
        run = api.add_field_run(paragraph, "PAGE")
        assert run is not None

    def test_add_field_paragraph_mixed(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        parts = [
            {"type": "text", "value": "文档 "},
            {"type": "field", "code": r"SEQ Figure \* ARABIC"},
            {"type": "text", "value": " - 说明"},
        ]
        paragraph = api.add_field_paragraph(container.subdoc, parts, BODY_STYLE)
        assert paragraph is not None

    def test_add_field_paragraph_invalid_type(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        parts = [{"type": "invalid", "value": "test"}]
        with pytest.raises(ValueError):
            api.add_field_paragraph(container.subdoc, parts)

    def test_add_figure_caption_auto(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        paragraph = api.add_figure_caption_auto(
            container.subdoc,
            "系统架构图",
            CAPTION_STYLE,
        )
        assert paragraph is not None
        xml = paragraph._element.xml
        assert r"STYLEREF KL一级标题 \n \* MERGEFORMAT " in xml
        assert r"SEQ 图 \* ARABIC \s 1" in xml
        assert "系统架构图" in xml


    def test_add_table_caption_auto(self, template_path):
        api = WordAPI(template_path)
        container = api.new_container()
        paragraph = api.add_table_caption_auto(
            container.subdoc,
            "数据统计表",
            CAPTION_STYLE,
        )
        assert paragraph is not None
        xml = paragraph._element.xml
        assert r"STYLEREF KL一级标题 \n \* MERGEFORMAT " in xml
        assert r"SEQ 表 \* ARABIC \s 1" in xml
        assert "数据统计表" in xml

    # =========================
    # 8. render 能力
    # =========================
    def test_render_basic(self, template_path, output_dir):
        api = WordAPI(template_path)
        output_file = str(output_dir / "test_render_basic.docx")
        result_path = api.render(build_min_context(api), output_file)
        assert os.path.exists(result_path)
        assert result_path == output_file

    def test_render_with_content(self, template_path, test_image, output_dir):
        api = WordAPI(template_path)
        output_file = str(output_dir / "test_render_with_content.docx")

        image_container = api.new_container()
        image_container.add_image(test_image, width_cm=6.0)

        table_container = api.new_container()
        table_container.add_table([
            ["项目", "数值"],
            ["指标1", "100"],
            ["指标2", "200"],
        ])

        result_container = api.new_container()
        result_container.add_heading("测试结果", level=1)
        result_container.add_paragraph("所有测试通过")

        context = {
            "text_tag": make_rich_text("自动化测试报告", BODY_STYLE),
            "image_tag": image_container.subdoc,
            "table_tag": table_container.subdoc,
            "result": result_container.subdoc,
        }

        result_path = api.render(context, output_file)
        assert os.path.exists(result_path)

    def test_render_creates_directory(self, template_path, tmp_path):
        api = WordAPI(template_path)
        output_file = str(tmp_path / "nested" / "dir" / "output.docx")
        result_path = api.render(build_min_context(api), output_file)
        assert os.path.exists(result_path)

    # =========================
    # 9. 页眉页脚（严格只基于 render 结果）
    # =========================
    def test_write_header_footer_after_render(self, template_path, output_dir):
        api = WordAPI(template_path)
        output_file = str(output_dir / "test_header_footer.docx")

        result_path = api.render(build_min_context(api), output_file)
        assert os.path.exists(result_path)

        api.write_header_footer(result_path, header_text="全局页眉")

        document = Document(result_path)
        assert len(document.sections) >= 1

        # 页眉：所有节都要有
        for section in document.sections:
            header_text = "".join(p.text for p in section.header.paragraphs)
            assert "全局页眉" in header_text

        # 页脚：最后一节必须有 PAGE / SECTIONPAGES
        footer_xml = document.sections[-1].footer.paragraphs[0]._element.xml
        assert "PAGE" in footer_xml
        assert "SECTIONPAGES" in footer_xml

        # 页码起始：最后一节从1开始
        sect_pr_xml = document.sections[-1]._sectPr.xml
        assert 'w:start="1"' in sect_pr_xml or "w:start='1'" in sect_pr_xml

    def test_render_only_has_no_real_footer(self, template_path, output_dir):
        api = WordAPI(template_path)
        output_file = str(output_dir / "test_render_only.docx")

        result_path = api.render(build_min_context(api), output_file)
        assert os.path.exists(result_path)

        document = Document(result_path)
        footer_xml = document.sections[-1].footer.paragraphs[0]._element.xml
        assert "SECTIONPAGES" not in footer_xml

    # =========================
    # 10. 辅助方法
    # =========================
    def test_check_color_valid(self):
        assert WordAPI._check_color("#FF0000") == "FF0000"
        assert WordAPI._check_color("00FF00") == "00FF00"
        assert WordAPI._check_color(None) is None

    def test_check_color_invalid(self):
        with pytest.raises(ValueError):
            WordAPI._check_color("#GGG")
        with pytest.raises(ValueError):
            WordAPI._check_color("12345")

    def test_get_paragraph_alignment(self):
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        assert WordAPI._get_paragraph_alignment("center") == WD_PARAGRAPH_ALIGNMENT.CENTER
        assert WordAPI._get_paragraph_alignment("right") == WD_PARAGRAPH_ALIGNMENT.RIGHT
        assert WordAPI._get_paragraph_alignment("left") == WD_PARAGRAPH_ALIGNMENT.LEFT
        assert WordAPI._get_paragraph_alignment(None) == WD_PARAGRAPH_ALIGNMENT.LEFT

    def test_get_vertical_alignment(self):
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        assert WordAPI._get_vertical_alignment("top") == WD_CELL_VERTICAL_ALIGNMENT.TOP
        assert WordAPI._get_vertical_alignment("bottom") == WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        assert WordAPI._get_vertical_alignment("center") == WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def test_get_table_alignment(self):
        from docx.enum.table import WD_TABLE_ALIGNMENT
        assert WordAPI._get_table_alignment("left") == WD_TABLE_ALIGNMENT.LEFT
        assert WordAPI._get_table_alignment("right") == WD_TABLE_ALIGNMENT.RIGHT
        assert WordAPI._get_table_alignment("center") == WD_TABLE_ALIGNMENT.CENTER

    def test_is_image_file(self, template_path):
        api = WordAPI(template_path)
        assert api._is_image_file("test.png") is True
        assert api._is_image_file("test.jpg") is True
        assert api._is_image_file("test.docx") is False

    def test_get_display_length(self, template_path):
        api = WordAPI(template_path)
        assert api._get_display_length("abc") == 3
        assert api._get_display_length("中文") >= 4
        assert api._get_display_length("") == 1

    def test_normalize_table_data(self, template_path):
        api = WordAPI(template_path)
        data = [["A", "B"], ["1"]]
        normalized = api._normalize_table_data(data)
        assert len(normalized[1]) == 2

    def test_scale_widths_to_max(self, template_path):
        api = WordAPI(template_path)
        widths = [5.0, 5.0, 5.0]
        scaled = api._scale_widths_to_max(widths, 10.0)
        assert sum(scaled) <= 10.0

    # =========================
    # 11. 集成测试
    # =========================
    def test_full_document_generation(self, template_path, test_image, output_dir):
        api = WordAPI(template_path)
        output_file = str(output_dir / "full_document.docx")

        text_content = make_rich_text("完整测试文档", BODY_STYLE)

        image_container = api.new_container()
        image_container.add_image(test_image, width_cm=8.0)
        image_container.add_figure_caption_auto("测试图片")

        table_container = api.new_container()
        table_container.add_heading("数据表格", level=2)
        table_container.add_table([
            ["序号", "项目名称", "状态"],
            ["1", "项目A", "完成"],
            ["2", "项目B", "进行中"],
        ])
        table_container.add_table_caption_auto("项目状态表")

        result_container = api.new_container()
        result_container.add_title("测试结论")
        result_container.add_paragraph("所有功能测试通过")
        result_container.add_page_break()
        result_container.add_paragraph("第二页内容")

        context = {
            "text_tag": text_content,
            "image_tag": image_container.subdoc,
            "table_tag": table_container.subdoc,
            "result": result_container.subdoc,
        }

        final_path = api.render(context, output_file)
        assert os.path.exists(final_path)
        assert os.path.getsize(final_path) > 0

    def test_generate_sample_document(self, template_path, test_image):
        output_dir = Path(__file__).parent / "output"
        output_dir.mkdir(exist_ok=True)
        output_file = str(output_dir / "sample_test_output.docx")

        api = WordAPI(template_path)
        text_content = make_rich_text("docxspec 完整功能测试报告", BODY_STYLE)

        image_container = api.new_container()
        image_container.add_title("图片功能测试")
        image_container.add_heading("测试项1：仅指定宽度插入图片", level=2)
        image_container.add_paragraph("测试方法：add_image(image_path, width_cm=6.0)")
        image_container.add_paragraph("测试结果：✓ 通过")
        image_container.add_paragraph("说明：图片按指定宽度等比例缩放")
        image_container.add_image(test_image, width_cm=6.0)
        image_container.add_figure_caption_auto("仅指定宽度")

        image_container.add_heading("测试项2：仅指定高度插入图片", level=2)
        image_container.add_paragraph("测试方法：add_image(image_path, height_cm=4.0)")
        image_container.add_paragraph("测试结果：✓ 通过")
        image_container.add_paragraph("说明：图片按指定高度等比例缩放")
        image_container.add_image(test_image, height_cm=4.0)
        image_container.add_figure_caption_auto("仅指定高度")

        image_container.add_heading("测试项3：同时指定宽度和高度", level=2)
        image_container.add_paragraph("测试方法：add_image(image_path, width_cm=10.0, height_cm=6.0)")
        image_container.add_paragraph("测试结果：✓ 通过")
        image_container.add_paragraph("说明：图片按指定尺寸强制缩放")
        image_container.add_image(test_image, width_cm=10.0, height_cm=6.0)
        image_container.add_figure_caption_auto("同时指定宽高")

        image_container.add_heading("测试项4：不指定尺寸使用原始大小", level=2)
        image_container.add_paragraph("测试方法：add_image(image_path)")
        image_container.add_paragraph("测试结果：✓ 通过")
        image_container.add_paragraph("说明：图片保持原始分辨率")
        image_container.add_image(test_image)
        image_container.add_figure_caption_auto("原始大小")

        table_container = api.new_container()
        table_container.add_title("表格功能测试")

        table_container.add_heading("测试项1：基础表格插入", level=2)
        table_container.add_paragraph("测试方法：insert_table(container, data)")
        table_container.add_paragraph("测试结果：✓ 通过")
        table_container.add_paragraph("说明：验证表格行数和列数正确")
        table_container.add_table([
            ["姓名", "年龄", "城市"],
            ["张三", "25", "北京"],
            ["李四", "30", "上海"],
        ])
        table_container.add_table_caption_auto("基础表格")

        table_container.add_heading("测试项2：自定义样式表格", level=2)
        table_container.add_paragraph("测试方法：insert_table(data, header_style, body_style, table_style)")
        table_container.add_paragraph("测试结果：✓ 通过")
        table_container.add_paragraph("说明：支持表头表体自定义样式")
        header_style = make_cell_style(bg_color="4472C4", font_color="FFFFFF", bold=True)
        body_style = make_cell_style(font_size=10)
        table_style = make_table_style(border_color="4472C4")
        table_container.add_table(
            [["项目", "状态"], ["测试A", "完成"], ["测试B", "进行中"]],
            header_style=header_style,
            body_style=body_style,
            table_style=table_style,
        )
        table_container.add_table_caption_auto("自定义样式表格")

        table_container.add_heading("测试项3：表格单元格插入图片", level=2)
        table_container.add_paragraph("测试方法：insert_table(data)，其中单元格值为图片路径")
        table_container.add_paragraph("测试结果：✓ 通过")
        table_container.add_paragraph("说明：当单元格值为有效图片路径时，自动插入图片")
        table_container.add_table([["图片", "说明"], [test_image, "表格中的图片"]])
        table_container.add_table_caption_auto("表格插图")

        result_container = api.new_container()
        result_container.add_title("其他功能测试")

        result_container.add_heading("段落功能", level=2)
        result_container.add_heading("测试项1：添加空段落", level=3)
        result_container.add_paragraph("测试方法：add_empty_paragraph(container, style)")
        result_container.add_paragraph("测试结果：✓ 通过")
        api.add_empty_paragraph(result_container.subdoc, BODY_STYLE)

        result_container.add_heading("测试项2：添加文本段(Run)", level=3)
        result_container.add_paragraph("测试方法：add_text_run(paragraph, text, style)")
        result_container.add_paragraph("测试结果：✓ 通过")
        para = result_container.subdoc.add_paragraph()
        api.add_text_run(para, "这是文本段测试", BODY_STYLE)

        result_container.add_heading("测试项3：添加完整段落", level=3)
        result_container.add_paragraph("测试方法：add_paragraph(container, text, style)")
        result_container.add_paragraph("测试结果：✓ 通过")
        api.add_paragraph(result_container.subdoc, "这是完整段落测试", BODY_STYLE)

        result_container.add_heading("域代码功能", level=2)
        result_container.add_heading("测试项1：添加单个域代码", level=3)
        result_container.add_paragraph("测试方法：add_field_run(paragraph, 'PAGE')")
        result_container.add_paragraph("测试结果：✓ 通过")
        para = result_container.subdoc.add_paragraph()
        api.add_text_run(para, "当前页码：", BODY_STYLE)
        api.add_field_run(para, "PAGE")

        result_container.add_heading("题注和页眉页脚", level=2)
        result_container.add_heading("测试项1：自动编号图注", level=3)
        result_container.add_paragraph("测试方法：add_figure_caption_auto('标题')")
        result_container.add_paragraph("测试结果：✓ 通过")
        api.add_figure_caption_auto(result_container.subdoc, "示例图注")

        result_container.add_heading("测试项2：自动编号表注", level=3)
        result_container.add_paragraph("测试方法：add_table_caption_auto('标题')")
        result_container.add_paragraph("测试结果：✓ 通过")
        api.add_table_caption_auto(result_container.subdoc, "示例表注")

        result_container.add_heading("测试项3：页眉页脚写入", level=3)
        result_container.add_paragraph("测试方法：render(output_path) + write_header_footer(output_path)")
        result_container.add_paragraph("测试结果：✓ 通过")
        result_container.add_paragraph("说明：完全基于 render 结果进行页眉页脚处理，不在测试里篡改结构")

        context = {
            "text_tag": text_content,
            "image_tag": image_container.subdoc,
            "table_tag": table_container.subdoc,
            "result": result_container.subdoc,
        }

        final_path = api.render(context, output_file)
        api.write_header_footer(final_path, header_text="示例页眉")

        print(f"\\n示例文档已生成: {final_path}")
        assert os.path.exists(final_path)
