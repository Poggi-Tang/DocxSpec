from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document

from docxspec import (
    mark_update_fields_on_open,
    normalize_table_list_fields,
    refresh_docx_fields,
)


def _replace_docx_part(path: Path, part_name: str, text: str) -> None:
    tmp_path = path.with_suffix(".tmp.docx")
    with ZipFile(path, "r") as src:
        with ZipFile(tmp_path, "w", ZIP_DEFLATED) as dst:
            for info in src.infolist():
                data = (
                    text.encode("utf-8")
                    if info.filename == part_name
                    else src.read(info.filename)
                )
                dst.writestr(info, data)
    tmp_path.replace(path)


def _read_docx_part(path: Path, part_name: str) -> str:
    with ZipFile(path, "r") as docx:
        return docx.read(part_name).decode("utf-8")


def test_mark_update_fields_on_open_sets_word_setting(tmp_path: Path) -> None:
    docx = tmp_path / "fields.docx"
    Document().save(docx)

    mark_update_fields_on_open(docx)
    settings = _read_docx_part(docx, "word/settings.xml")

    assert "updateFields" in settings
    assert 'w:val="true"' in settings


def test_normalize_table_list_fields_uses_standard_table_label(tmp_path: Path) -> None:
    docx = tmp_path / "table_list.docx"
    doc = Document()
    doc.add_paragraph("表目录")
    doc.save(docx)

    document_xml = _read_docx_part(docx, "word/document.xml")
    document_xml = document_xml.replace(
        "表目录",
        r'TOC \h \c "表格"',
    )
    _replace_docx_part(docx, "word/document.xml", document_xml)

    count = normalize_table_list_fields(docx)
    normalized_xml = _read_docx_part(docx, "word/document.xml")

    assert count == 1
    assert r'TOC \h \z \c "表"' in normalized_xml
    assert r'TOC \h \c "表格"' not in normalized_xml


def test_refresh_docx_fields_can_skip_word_com(tmp_path: Path) -> None:
    docx = tmp_path / "refresh.docx"
    doc = Document()
    doc.add_paragraph("表目录")
    doc.save(docx)

    document_xml = _read_docx_part(docx, "word/document.xml")
    _replace_docx_part(
        docx,
        "word/document.xml",
        document_xml.replace("表目录", r'TOC \h \z \c "表格"'),
    )

    result = refresh_docx_fields(docx, use_word=False)
    settings = _read_docx_part(docx, "word/settings.xml")
    normalized_xml = _read_docx_part(docx, "word/document.xml")

    assert result.update_fields_on_open is True
    assert result.table_list_fields_normalized == 1
    assert result.word_refreshed is False
    assert result.error is None
    assert "updateFields" in settings
    assert r'TOC \h \z \c "表"' in normalized_xml
