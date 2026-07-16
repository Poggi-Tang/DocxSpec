"""KL 标准能力测试。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from docxspec import (
    DEFAULT_MASTER_TEMPLATE,
    build_standard_docx,
    check_word_standard,
    classify_docx_body,
    classify_heading,
    clean_heading_number,
    detect_source_front_matter,
)


def test_clean_heading_number_and_classify_heading() -> None:
    cleaned, kind = clean_heading_number("1.2 接口设计")

    assert cleaned == "接口设计"
    assert kind == "heading_2"
    assert clean_heading_number("1.2.3.4 四级标题")[1] == "heading_4"
    assert clean_heading_number("1.2.3.4.5 五级标题")[1] == "heading_5"
    assert clean_heading_number("1.2.3.4.5.6 六级标题")[1] == "heading_6"
    assert classify_heading("第一章 需求背景")[0] == "heading_1"
    assert classify_heading("① 这是正文条目")[0] == "body"


def test_classify_docx_body_reports_headings_tables_and_captions(tmp_path: Path) -> None:
    docx = tmp_path / "body.docx"
    doc = Document()
    doc.add_paragraph("一、需求背景")
    doc.add_paragraph("这是正文内容。")
    doc.add_paragraph("表 1-1 数据表")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "名称"
    table.cell(1, 1).text = "积分元件"
    doc.save(docx)

    report = classify_docx_body(docx)

    assert report["summary"]["heading_1"] == 1
    assert report["summary"]["body"] == 1
    assert report["summary"]["table_caption"] == 1
    assert report["summary"]["table"] == 1
    assert report["tables"][0]["caption"]["position"] == "above"


def test_list_paragraph_plain_text_is_not_heading(tmp_path: Path) -> None:
    docx = tmp_path / "list_paragraph.docx"
    doc = Document()
    doc.add_paragraph("概述", style="Heading 1")
    doc.add_paragraph("元件需求文档编制人员提出", style="List Paragraph")
    doc.save(docx)

    report = classify_docx_body(docx)

    assert report["blocks"][1]["type"] == "body"


def test_code_brace_lines_are_not_headings(tmp_path: Path) -> None:
    docx = tmp_path / "code_braces.docx"
    output = tmp_path / "code_braces_standard.docx"
    doc = Document()
    doc.add_paragraph("伪代码：")
    doc.add_paragraph("{")
    doc.add_paragraph("model_cv.Integrator_CState = model_p.Integrator_IC;")
    doc.add_paragraph("}")
    doc.save(docx)

    source_report = classify_docx_body(docx)
    assert [block["type"] for block in source_report["blocks"]] == [
        "body",
        "body",
        "body",
        "body",
    ]

    build_standard_docx(docx, output)
    output_report = classify_docx_body(output)
    brace_blocks = [
        block for block in output_report["blocks"] if block.get("text", "").strip() in {"{", "}"}
    ]

    assert brace_blocks
    assert all(block["type"] == "body" for block in brace_blocks)


def test_code_like_lines_are_not_headings(tmp_path: Path) -> None:
    docx = tmp_path / "code_lines.docx"
    output = tmp_path / "code_lines_standard.docx"
    code_lines = [
        "if (IsMajorTimeStep(model)) {",
        "local_stateReduction(&model_cv.Integrator_CState, 1， &PeriodicContStateRanges);",
        "model_cv.Integrator_CState = model_p.Integrator_IC;",
        "} else if (model_cv.Integrator_CState <= LowerSa) {",
        "|| (usat && (u_1 < 0.0))) {",
        "Xdot->Integrator_CState = u_1;",
        "} else {",
        "/* in saturation */",
        "_rtXdot->Integrator_CState = 0.0;",
    ]

    doc = Document()
    for line in code_lines:
        doc.add_paragraph(line)
    doc.save(docx)

    source_report = classify_docx_body(docx)
    assert [block["type"] for block in source_report["blocks"]] == ["body"] * len(code_lines)

    build_standard_docx(docx, output)
    output_report = classify_docx_body(output)
    matched = [
        block
        for block in output_report["blocks"]
        if block.get("text", "").strip() in code_lines
    ]

    assert len(matched) == len(code_lines)
    assert all(block["type"] == "body" for block in matched)


def test_single_level_port_items_are_not_promoted_to_headings(tmp_path: Path) -> None:
    docx = tmp_path / "port_items.docx"
    output = tmp_path / "port_items_standard.docx"
    doc = Document()
    doc.add_paragraph("3.1.1.2 输入")
    doc.add_paragraph("该模块具备三种类型的输入信号。")
    doc.add_paragraph("1）端口1-被积信号")
    doc.add_paragraph("被积信号是被积分的信号源。")
    doc.add_paragraph("2）端口2-外部重置信号")
    doc.add_paragraph("外部重置信号仅支持输入Double、Boolean类型的实数。")
    doc.save(docx)

    source_report = classify_docx_body(docx)
    by_text = {block["text"]: block["type"] for block in source_report["blocks"]}

    assert by_text["3.1.1.2 输入"] == "heading_4"
    assert by_text["1）端口1-被积信号"] == "body"
    assert by_text["2）端口2-外部重置信号"] == "body"

    build_standard_docx(docx, output)
    output_report = classify_docx_body(output)
    output_by_text = {block["text"]: block["type"] for block in output_report["blocks"]}

    assert output_by_text["输入"] == "heading_4"
    assert output_by_text["1）端口1-被积信号"] == "body"
    assert output_by_text["2）端口2-外部重置信号"] == "body"


def test_short_plain_text_without_heading_style_or_number_is_body(tmp_path: Path) -> None:
    docx = tmp_path / "plain_short_text.docx"
    output = tmp_path / "plain_short_text_standard.docx"
    doc = Document()
    doc.add_paragraph("3.5.1.1 调研渠道")
    doc.add_paragraph("Powerfactory 2026软件")
    doc.save(docx)

    source_report = classify_docx_body(docx)
    by_text = {block["text"]: block["type"] for block in source_report["blocks"]}

    assert by_text["3.5.1.1 调研渠道"] == "heading_4"
    assert by_text["Powerfactory 2026软件"] == "body"

    build_standard_docx(docx, output)
    output_report = classify_docx_body(output)
    output_by_text = {block["text"]: block["type"] for block in output_report["blocks"]}

    assert output_by_text["调研渠道"] == "heading_4"
    assert output_by_text["Powerfactory 2026软件"] == "body"


def test_blank_line_is_inserted_before_headings_after_body_only(tmp_path: Path) -> None:
    docx = tmp_path / "heading_spacing.docx"
    output = tmp_path / "heading_spacing_standard.docx"
    doc = Document()
    doc.add_paragraph("1 一级标题")
    doc.add_paragraph("正文段落")
    doc.add_paragraph("1.1 二级标题")
    doc.add_paragraph("1.1.1 三级标题")
    doc.save(docx)

    build_standard_docx(docx, output)
    paragraphs = [paragraph.text.strip() for paragraph in Document(str(output)).paragraphs]
    h2_index = paragraphs.index("二级标题")
    h3_index = paragraphs.index("三级标题")

    assert paragraphs[h2_index - 1] == ""
    assert paragraphs[h3_index - 1] == "二级标题"


def test_check_word_standard_uses_template_styles() -> None:
    template = Path(__file__).parent / "templates" / "template.docx"

    report = check_word_standard(template)

    assert report["required_styles_ok"] is True
    assert report["missing_styles"] == []


def test_default_master_template_is_packaged() -> None:
    assert DEFAULT_MASTER_TEMPLATE.exists()
    report = check_word_standard(DEFAULT_MASTER_TEMPLATE)
    assert report["required_styles_ok"] is True
    style_names = report["found_styles"]
    assert "KL四级标题" in style_names
    assert "KL五级标题" in style_names
    assert "KL六级标题" in style_names
    assert "KL其他标题" not in style_names


def test_build_standard_docx_from_body_document(tmp_path: Path) -> None:
    template = Path(__file__).parent / "templates" / "template.docx"
    body = tmp_path / "body.docx"
    output = tmp_path / "standard.docx"

    doc = Document()
    doc.add_paragraph("积分元件测试大纲")
    doc.add_paragraph("1 需求背景")
    doc.add_paragraph("这是正文内容。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "名称"
    table.cell(1, 1).text = "积分元件"
    doc.save(body)

    report = build_standard_docx(body, output, template)

    assert output.exists()
    assert report["source_title"] == "积分元件测试大纲"
    assert report["tables_styled"] >= 1
    assert report["table_captions_added"]
    standard_report = check_word_standard(output)
    assert standard_report["required_styles_ok"] is True
    assert standard_report["has_table_seq"] is True


def test_existing_caption_text_is_rebuilt_with_standard_fields(tmp_path: Path) -> None:
    body = tmp_path / "caption_body.docx"
    output = tmp_path / "caption_standard.docx"

    doc = Document()
    doc.add_paragraph("1 需求背景")
    doc.add_paragraph("表 9-8 非标准表题注")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "名称"
    table.cell(1, 1).text = "积分元件"
    doc.save(body)

    report = build_standard_docx(body, output)
    result = Document(str(output))
    caption = next(paragraph for paragraph in result.paragraphs if "非标准表题注" in paragraph.text)
    xml = caption._p.xml

    assert report["captions_standardized"]
    assert "STYLEREF KL一级标题" in xml
    assert r"SEQ 表 \* ARABIC \r 1" in xml
    assert "非标准表题注" in caption.text
    assert len([paragraph for paragraph in result.paragraphs if "非标准表题注" in paragraph.text]) == 1


def test_caption_sequences_explicitly_reset_for_each_chapter(tmp_path: Path) -> None:
    body = tmp_path / "chapter_tables.docx"
    output = tmp_path / "chapter_tables_standard.docx"

    doc = Document()
    for chapter in ("第一章", "第二章"):
        doc.add_paragraph(chapter, style="Heading 1")
        for index in (1, 2):
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "章节"
            table.cell(0, 1).text = "序号"
            table.cell(1, 0).text = chapter
            table.cell(1, 1).text = str(index)
    doc.save(body)

    report = build_standard_docx(body, output)
    result = Document(str(output))
    captions = [
        paragraph
        for paragraph in result.paragraphs
        if "SEQ 表" in paragraph._p.xml
    ][-4:]
    field_xml = [paragraph._p.xml for paragraph in captions]

    assert len(report["table_captions_added"]) == 4
    assert len(captions) == 4
    assert r"SEQ 表 \* ARABIC \r 1" in field_xml[0]
    assert r"SEQ 表 \* ARABIC \r 1" not in field_xml[1]
    assert r"SEQ 表 \* ARABIC \r 1" in field_xml[2]
    assert r"SEQ 表 \* ARABIC \r 1" not in field_xml[3]


def test_build_standard_docx_uses_packaged_template_by_default(tmp_path: Path) -> None:
    body = tmp_path / "body.docx"
    output = tmp_path / "standard_default.docx"

    doc = Document()
    doc.add_paragraph("积分元件测试大纲")
    doc.add_paragraph("1 需求背景")
    doc.add_paragraph("这是正文内容。")
    doc.save(body)

    report = build_standard_docx(body, output)

    assert output.exists()
    assert report["template"] == str(DEFAULT_MASTER_TEMPLATE)
    assert check_word_standard(output)["required_styles_ok"] is True


def test_build_standard_docx_preserves_deep_heading_levels(tmp_path: Path) -> None:
    body = tmp_path / "deep_headings.docx"
    output = tmp_path / "deep_standard.docx"

    doc = Document()
    doc.add_paragraph("积分元件测试大纲")
    for text in [
        "1 一级标题",
        "1.1 二级标题",
        "1.1.1 三级标题",
        "1.1.1.1 四级标题",
        "1.1.1.1.1 五级标题",
        "1.1.1.1.1.1 六级标题",
    ]:
        doc.add_paragraph(text)
    doc.save(body)

    build_standard_docx(body, output)
    report = classify_docx_body(output)
    summary = report["summary"]

    assert summary["heading_1"] >= 1
    assert summary["heading_2"] >= 1
    assert summary["heading_3"] >= 1
    assert summary["heading_4"] >= 1
    assert summary["heading_5"] >= 1
    assert summary["heading_6"] >= 1


def test_build_standard_docx_preserves_source_heading_styles_after_compose(tmp_path: Path) -> None:
    body = tmp_path / "source_style_headings.docx"
    output = tmp_path / "source_style_standard.docx"

    doc = Document()
    doc.add_paragraph("积分元件测试大纲")
    for level in range(1, 7):
        style_name = f"Heading {level}"
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        doc.add_paragraph(f"{level}级样式标题", style=style_name)
    doc.save(body)

    build_standard_docx(body, output)
    report = classify_docx_body(output)
    summary = report["summary"]

    assert summary["heading_1"] >= 1
    assert summary["heading_2"] >= 1
    assert summary["heading_3"] >= 1
    assert summary["heading_4"] >= 1
    assert summary["heading_5"] >= 1
    assert summary["heading_6"] >= 1


def test_detect_source_front_matter_does_not_skip_direct_body(tmp_path: Path) -> None:
    body = tmp_path / "direct_body.docx"

    doc = Document()
    doc.add_paragraph("1 需求背景")
    doc.add_paragraph("这是正文内容。")
    doc.save(body)

    source_report = classify_docx_body(body)
    front_matter = detect_source_front_matter(source_report)

    assert front_matter["detected"] is False
    assert front_matter["body_start_block"] == 0
    assert front_matter["skipped_blocks_count"] == 0


def test_build_standard_docx_keeps_first_heading_when_body_starts_immediately(tmp_path: Path) -> None:
    body = tmp_path / "积分器需求文档.docx"
    output = tmp_path / "direct_body_standard.docx"

    doc = Document()
    doc.add_paragraph("概述", style="Heading 1")
    doc.add_paragraph("这是正文内容。")
    doc.save(body)

    report = build_standard_docx(body, output)
    output_doc = Document(str(output))
    paragraphs = [paragraph.text for paragraph in output_doc.paragraphs]

    assert report["source_title"] == "积分器需求文档"
    assert "概述" in paragraphs


def test_build_standard_docx_preserves_source_from_first_block_by_default(tmp_path: Path) -> None:
    body = tmp_path / "front_matter_like_body.docx"
    output = tmp_path / "front_matter_like_body_standard.docx"

    doc = Document()
    doc.add_paragraph("文件编号：KL-JS-001")
    doc.add_paragraph("概述", style="Heading 1")
    doc.add_paragraph("正文从这里开始。")
    doc.save(body)

    report = build_standard_docx(body, output)
    output_text = "\n".join(paragraph.text for paragraph in Document(str(output)).paragraphs)

    assert report["skip_source_front_matter"] is False
    assert report["source_front_matter"]["body_start_block"] == 0
    assert "文件编号：KL-JS-001" in output_text
    assert "概述" in output_text


def test_build_standard_docx_can_skip_source_front_matter_when_requested(tmp_path: Path) -> None:
    body = tmp_path / "front_matter.docx"
    output = tmp_path / "front_matter_standard.docx"

    doc = Document()
    doc.add_paragraph("文件编号：KL-JS-001")
    doc.add_paragraph("SimuNPS控制元件")
    doc.add_paragraph("开发和测试需求文档")
    doc.add_paragraph("上海科梁信息科技股份有限公司")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "版本"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "V1.0"
    table.cell(1, 1).text = "创建"
    doc.add_paragraph("目录")
    doc.add_paragraph("概述", style="Heading 1")
    doc.add_paragraph("正文从这里开始。")
    doc.add_paragraph("调研渠道", style="Heading 4")
    doc.add_paragraph("这是四级标题下的正文。")
    doc.save(body)

    report = build_standard_docx(body, output, skip_source_front_matter=True)
    output_doc = Document(str(output))
    paragraphs = [paragraph.text for paragraph in output_doc.paragraphs]
    output_text = "\n".join(paragraphs)
    output_report = classify_docx_body(output)

    assert report["source_title"] == "SimuNPS控制元件开发和测试需求文档"
    assert report["source_file_number"] == "文件编号：KL-JS-001"
    assert report["source_front_matter"]["detected"] is True
    assert report["source_front_matter"]["body_start_block"] > 0
    assert "文件编号：KL-JS-001" in output_text
    assert "KL-JS-PCXX-XX" not in output_text
    assert "创建" not in output_text
    assert "图注1" not in output_text
    assert "表注1" not in output_text
    assert output_report["summary"]["heading_1"] >= 1
    assert output_report["summary"]["heading_4"] >= 1
    assert check_word_standard(output)["figure_table_lists_ok"] is True
