# -*- coding: utf-8 -*-
"""Word 文档生成 API。

本模块基于 ``docxtpl`` 与 ``python-docx`` 提供统一的报告生成接口：

1. ``docxtpl`` 负责模板占位渲染。
2. ``python-docx`` 负责段落、表格、图片、页眉页脚及域的补充写入。
3. 通过 :class:`DocContainer` 提供链式子文档拼装能力。
"""

from __future__ import annotations

import copy
import os
import unicodedata
from pathlib import Path
from typing import Any, Optional, Sequence, Union

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docxtpl import DocxTemplate
from PIL import Image

from .word_styles import (
    BODY_STYLE,
    CAPTION_STYLE,
    FOOTER_STYLE,
    HEADER_STYLE,
    IMAGE_STYLE,
    TABLE_BODY_STYLE,
    TABLE_HEADER_STYLE,
    CellStyle,
    TableStyle,
    TextStyle,
    make_cell_style,
    make_table_style,
    make_text_style,
)

TextValue = Union[str, int, float, None]
PartConfig = dict[str, Any]
TableConfig = dict[str, Any]


class BlockTemplate:
    """从 Word 模板中抽取出来、可重复使用的 OpenXML 块。

    该对象直接保存 WordprocessingML 原始节点，而不是重新创建段落或表格。
    因此，块中的题注域、表格样式、合并单元格、图片、段落样式等底层格式
    在 clone 后再次插入文档时可以尽量保持不变。
    """

    def __init__(self, elements: Sequence[Any]) -> None:
        """保存块节点的深拷贝，避免后续修改污染来源模板。"""
        if not elements:
            raise ValueError("块模板不能为空")
        self.elements = [copy.deepcopy(element) for element in elements]

    def clone(self) -> "BlockTemplate":
        """复制一个新的块实例，用于同一模板块的多次填充和插入。"""
        return BlockTemplate(self.elements)

    def replace_text(self, mapping: dict[str, Any]) -> "BlockTemplate":
        """替换块内所有 ``w:t`` 文本节点中的占位符。

        这里有意只替换文本节点，不重建 run 或 paragraph。这样可以保留
        占位符周围原有的字体、域代码、题注结构等格式信息。需要注意的是，
        该方法适用于占位符没有被 Word 拆分到多个 run 的常见场景。
        """
        if not mapping:
            return self

        normalized = {
            str(key): "" if value is None else str(value)
            for key, value in mapping.items()
        }
        for element in self.elements:
            for text_node in element.xpath(".//w:t"):
                if text_node.text is None:
                    continue
                value = text_node.text
                for old, new in normalized.items():
                    value = value.replace(old, new)
                text_node.text = value
        return self

    def table(self, index: int = 0) -> Any:
        """返回块内第 ``index`` 个表格的原始 XML 节点。"""
        tables = [
            element
            for element in self.elements
            if getattr(element, "tag", None) == qn("w:tbl")
        ]
        try:
            return tables[index]
        except IndexError as exc:
            raise IndexError(f"块模板中不存在第 {index} 个表格") from exc


class DocContainer:
    """子文档容器，用于链式调用构建 Word 文档内容。

    该类封装了常用的文档元素添加操作，支持流畅的链式调用风格。

    :param api: WordAPI 实例引用
    :param subdoc: 子文档对象
    """

    def __init__(self, api: "WordAPI", subdoc: Any) -> None:
        self.api = api
        self.subdoc = subdoc

    def add_block(self, block: BlockTemplate) -> "DocContainer":
        """把一个块模板追加到容器中。

        该方法用于把已经从 Word 模板中抽取出来的题注、表格或其他连续节点
        放入 docxtpl 的子文档容器。块会被深拷贝后追加，调用方可以继续复用
        原始 ``BlockTemplate``。
        """
        self.api.append_block(self.subdoc, block)
        return self

    def add_title(self, text: str, style: Optional[TextStyle] = None) -> "DocContainer":
        """添加主标题。

        :param text: 标题文本
        :param style: 可选的文本样式，默认使用 KL主标题 样式
        :return: 当前容器实例，支持链式调用
        """
        self.api.add_paragraph(
            self.subdoc,
            text,
            style or make_text_style(style_name="KL主标题"),
        )
        return self

    def add_heading(
        self,
        text: str,
        level: int = 1,
        style: Optional[TextStyle] = None,
    ) -> "DocContainer":
        """添加分级标题。

        :param text: 标题文本
        :param level: 标题级别（1-3），默认为 1
        :param style: 可选的自定义样式，未指定时根据 level 自动选择
        :return: 当前容器实例，支持链式调用
        """
        if style is None:
            style = {
                1: make_text_style(style_name="KL一级标题"),
                2: make_text_style(style_name="KL二级标题"),
                3: make_text_style(style_name="KL三级标题"),
                4: make_text_style(style_name="KL四级标题"),
                5: make_text_style(style_name="KL五级标题"),
                6: make_text_style(style_name="KL六级标题"),
            }.get(level, BODY_STYLE)
        self.api.add_paragraph(self.subdoc, text, style)
        return self

    def add_paragraph(
        self,
        text: TextValue,
        style: Optional[TextStyle] = None,
    ) -> "DocContainer":
        """添加普通段落。

        :param text: 段落文本内容
        :param style: 可选的文本样式，默认使用正文样式
        :return: 当前容器实例，支持链式调用
        """
        self.api.add_paragraph(self.subdoc, text, style or BODY_STYLE)
        return self

    def add_image(
        self,
        image_path: str,
        width_cm: Optional[float] = None,
        height_cm: Optional[float] = None,
        align: str = "center",
        style: Optional[TextStyle] = None,
    ) -> "DocContainer":
        """添加图片块。

        :param image_path: 图片文件路径
        :param width_cm: 图片宽度（厘米），可选
        :param height_cm: 图片高度（厘米），可选
        :param align: 对齐方式（left/center/right），默认为 center
        :param style: 可选的图片样式，默认使用 IMAGE_STYLE
        :return: 当前容器实例，支持链式调用
        """
        self.api.add_image_block(
            self.subdoc,
            image_path=image_path,
            width_cm=width_cm,
            height_cm=height_cm,
            align=align,
            style=style or IMAGE_STYLE,
        )
        return self

    def add_table(
        self,
        data: list[list[TextValue]],
        header_style: Optional[CellStyle] = None,
        body_style: Optional[CellStyle] = None,
        table_style: Optional[TableStyle] = None,
    ) -> "DocContainer":
        """添加表格。

        :param data: 表格数据，二维列表结构
        :param header_style: 表头单元格样式，可选
        :param body_style: 表体单元格样式，可选
        :param table_style: 表格整体样式，可选
        :return: 当前容器实例，支持链式调用
        """
        self.api.insert_table(
            container=self.subdoc,
            data=data,
            header_style=header_style,
            body_style=body_style,
            table_style=table_style,
        )
        return self

    def add_table_by_config(self, table_config: TableConfig) -> "DocContainer":
        """通过配置字典添加表格。

        :param table_config: 表格配置字典，包含 data、style 等键
        :return: 当前容器实例，支持链式调用
        """
        self.api.insert_table_by_config(self.subdoc, table_config)
        return self

    def add_page_break(self) -> "DocContainer":
        """添加分页符。

        :return: 当前容器实例，支持链式调用
        """
        self.subdoc.add_page_break()
        return self

    def add_field_paragraph(
        self,
        parts: list[PartConfig],
        style: Optional[TextStyle] = None,
    ) -> Any:
        """添加包含域代码的段落。

        :param parts: 段落组成部分列表，每个部分为包含 type 和 value/code 的字典
        :param style: 段落样式，可选
        :return: 创建的段落对象
        """
        return self.api.add_field_paragraph(self.subdoc, parts, style or BODY_STYLE)

    def add_page_footer(self, style: Optional[TextStyle] = None) -> Any:
        """添加页脚页码信息。

        :param style: 页脚样式，可选
        :return: 创建的段落对象
        """
        return self.api.add_page_footer(self.subdoc, style or FOOTER_STYLE)

    def add_figure_caption_auto(
        self,
        title: str,
        style: Optional[TextStyle] = None,
    ) -> Any:
        """添加自动编号的图注。

        :param title: 图注标题文本
        :param style: 题注样式，可选
        :return: 创建的段落对象
        """
        return self.api.add_figure_caption_auto(
            self.subdoc,
            title,
            style or CAPTION_STYLE,
        )

    def add_table_caption_auto(
        self,
        title: str,
        style: Optional[TextStyle] = None,
    ) -> Any:
        """添加自动编号的表注。

        :param title: 表注标题文本
        :param style: 题注样式，可选
        :return: 创建的段落对象
        """
        return self.api.add_table_caption_auto(
            self.subdoc,
            title,
            style or CAPTION_STYLE,
        )


class WordAPI:
    """Word 报告生成主入口类。

    该类提供了完整的 Word 文档生成能力，包括：

    * 基于模板的文档渲染
    * 段落、表格、图片的插入
    * 域代码的支持（页码、图表编号等）
    * 页眉页脚的设置
    * 子文档容器的管理

    :param template_path: Word 模板文件路径
    :raises FileNotFoundError: 当模板文件不存在时抛出
    """

    IMAGE_EXTENSIONS = {
        ".png",
        ".jpg",
        ".jpeg",
        ".bmp",
        ".gif",
        ".webp",
        ".tif",
        ".tiff",
    }

    def __init__(self, template_path: str) -> None:
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板不存在: {template_path}")
        self.template_path = str(template_path)
        self.doc = DocxTemplate(self.template_path)

    def new_container(self) -> DocContainer:
        """创建新的子文档容器。

        :return: 新的 DocContainer 实例
        """
        return DocContainer(self, self.doc.new_subdoc())

    @staticmethod
    def _paragraph_text(paragraph_element: Any) -> str:
        return "".join(node.text or "" for node in paragraph_element.xpath(".//w:t"))

    @staticmethod
    def _is_paragraph(element: Any) -> bool:
        return getattr(element, "tag", None) == qn("w:p")

    @staticmethod
    def _is_table(element: Any) -> bool:
        return getattr(element, "tag", None) == qn("w:tbl")

    def _iter_body_elements(self) -> list[Any]:
        document = self.doc.get_docx()
        return list(document.element.body)

    def _find_marker_element(self, marker: str) -> Any:
        for element in self._iter_body_elements():
            if self._is_paragraph(element) and marker in self._paragraph_text(element):
                return element
        raise ValueError(f"未找到块标记: {marker}")

    def extract_table_block(
        self,
        marker: str,
        *,
        include_blank_before: bool = True,
        remove_marker: bool = False,
        remove_block: bool = False,
    ) -> BlockTemplate:
        """抽取标记段落后方的“题注 + 表格”块。

        典型用法是在 Word 模板中把 ``marker`` 放在标准表题注前一段。
        本方法会从该标记之后开始收集空段落、题注段落和紧随其后的表格，
        并直接保存原始 XML，因此可保留题注域和表格内部复杂格式。

        :param marker: 用于定位来源块的标记文本
        :param include_blank_before: 是否把标记和题注之间的空段落一并抽取
        :param remove_marker: 抽取后是否删除标记段落
        :param remove_block: 抽取后是否删除标记段落和来源块本身
        :return: 可复制、可填充、可插入的 :class:`BlockTemplate`
        """
        marker_element = self._find_marker_element(marker)
        elements = self._iter_body_elements()
        marker_index = elements.index(marker_element)

        collected: list[Any] = []
        seen_caption = False
        seen_table = False

        for element in elements[marker_index + 1:]:
            if self._is_paragraph(element):
                paragraph_text = self._paragraph_text(element).strip()
                if not paragraph_text and include_blank_before and not seen_caption:
                    collected.append(element)
                    continue
                if not seen_caption:
                    collected.append(element)
                    seen_caption = True
                    continue
                if not seen_table:
                    collected.append(element)
                    continue
                break

            if self._is_table(element):
                collected.append(element)
                seen_table = True
                break

            if seen_caption or collected:
                collected.append(element)

        if not seen_caption or not seen_table:
            raise ValueError(f"标记 {marker!r} 后未找到完整的题注+表格块")

        if remove_block:
            parent = marker_element.getparent()
            for element in [marker_element, *collected]:
                if element.getparent() is parent:
                    parent.remove(element)
        elif remove_marker:
            marker_element.getparent().remove(marker_element)

        return BlockTemplate(collected)

    def insert_block_at_marker(
        self,
        marker: str,
        block: BlockTemplate,
        *,
        remove_marker: bool = True,
    ) -> None:
        """把块插入到正文中指定标记段落之前。

        该方法绕过 ``docxtpl.render``，直接操作主文档 XML。适合需要尽量
        保留 Word 域题注、已有表格样式、图片关系等格式的场景。
        """
        marker_element = self._find_marker_element(marker)
        parent = marker_element.getparent()
        insert_index = parent.index(marker_element)

        for element in block.elements:
            parent.insert(insert_index, copy.deepcopy(element))
            insert_index += 1

        if remove_marker:
            parent.remove(marker_element)

    def append_block(self, container: Any, block: BlockTemplate) -> Any:
        """把块追加到子文档或类文档容器中。"""
        target = getattr(container, "_element", None)
        if target is None:
            target = getattr(container, "element", None)
        if target is None:
            raise TypeError("container 不支持追加 OpenXML 块")

        body = target.body if hasattr(target, "body") else target
        for element in block.elements:
            body.append(copy.deepcopy(element))
        return container

    @staticmethod
    def _check_color(color: Optional[str]) -> Optional[str]:
        """验证并标准化颜色值。

        :param color: 颜色值字符串（支持 #RRGGBB 或 RRGGBB 格式）
        :return: 标准化的 6 位大写十六进制颜色字符串
        :raises ValueError: 当颜色格式不正确时抛出
        """
        if color is None:
            return None
        value = color.strip().replace("#", "").upper()
        if len(value) != 6:
            raise ValueError(
                "颜色值必须为 6 位十六进制，例如 FF0000，"
                f"当前为: {color}"
            )
        return value

    @staticmethod
    def _get_paragraph_alignment(align: Optional[str]) -> WD_PARAGRAPH_ALIGNMENT:
        """获取段落对齐方式枚举值。

        :param align: 对齐方式字符串（left/center/right）
        :return: python-docx 的段落对齐枚举值
        """
        value = (align or "left").lower()
        if value == "center":
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        if value == "right":
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        return WD_PARAGRAPH_ALIGNMENT.LEFT

    @staticmethod
    def _get_vertical_alignment(
        vertical_align: Optional[str],
    ) -> WD_CELL_VERTICAL_ALIGNMENT:
        """获取单元格垂直对齐方式枚举值。

        :param vertical_align: 垂直对齐方式字符串（top/center/bottom）
        :return: python-docx 的单元格垂直对齐枚举值
        """
        value = (vertical_align or "center").lower()
        if value == "top":
            return WD_CELL_VERTICAL_ALIGNMENT.TOP
        if value == "bottom":
            return WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        return WD_CELL_VERTICAL_ALIGNMENT.CENTER

    @staticmethod
    def _get_table_alignment(align: Optional[str]) -> WD_TABLE_ALIGNMENT:
        """获取表格对齐方式枚举值。

        :param align: 对齐方式字符串（left/center/right）
        :return: python-docx 的表格对齐枚举值
        """
        value = (align or "center").lower()
        if value == "left":
            return WD_TABLE_ALIGNMENT.LEFT
        if value == "right":
            return WD_TABLE_ALIGNMENT.RIGHT
        return WD_TABLE_ALIGNMENT.CENTER

    @staticmethod
    def _set_run_font(
        run: Any,
        font_name: Optional[str] = None,
        font_size: Optional[float] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        font_color: Optional[str] = None,
    ) -> None:
        """设置文本段的字体属性。

        :param run: python-docx 的 Run 对象
        :param font_name: 字体名称
        :param font_size: 字号（磅）
        :param bold: 是否加粗
        :param italic: 是否斜体
        :param font_color: 字体颜色（十六进制格式）
        """
        font_color = WordAPI._check_color(font_color)

        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if font_name:
            run.font.name = font_name
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), font_name)
            rfonts.set(qn("w:ascii"), font_name)
            rfonts.set(qn("w:hAnsi"), font_name)
        if font_size is not None:
            run.font.size = Pt(font_size)
        if font_color:
            run.font.color.rgb = RGBColor.from_string(font_color)

    @staticmethod
    def _apply_paragraph_style(paragraph: Any, style_name: Optional[str]) -> None:
        """应用段落样式名称。

        :param paragraph: python-docx 的 Paragraph 对象
        :param style_name: 样式名称
        """
        if not style_name:
            return
        try:
            paragraph.style = style_name
        except Exception:
            return

    @staticmethod
    def _apply_paragraph_direct_format(
        paragraph: Any,
        style: Optional[TextStyle],
    ) -> None:
        """直接格式化段落属性。

        :param paragraph: python-docx 的 Paragraph 对象
        :param style: 文本样式配置对象
        """
        if style is None:
            return

        if style.align is not None:
            paragraph.alignment = WordAPI._get_paragraph_alignment(style.align)

        paragraph_format = paragraph.paragraph_format
        if style.line_spacing is not None:
            paragraph_format.line_spacing = style.line_spacing
            paragraph_format.line_spacing_rule = (
                WD_LINE_SPACING.SINGLE
                if style.line_spacing == 1
                else WD_LINE_SPACING.MULTIPLE
            )
        if style.space_before_pt is not None:
            paragraph_format.space_before = Pt(style.space_before_pt)
        if style.space_after_pt is not None:
            paragraph_format.space_after = Pt(style.space_after_pt)
        if style.first_line_indent_chars is not None:
            base_pt = style.font_size or 12.0
            paragraph_format.first_line_indent = Pt(
                base_pt * style.first_line_indent_chars
            )

    @staticmethod
    def _set_cell_background(cell: Any, fill: Optional[str]) -> None:
        """设置单元格背景色。

        :param cell: python-docx 的 Cell 对象
        :param fill: 填充颜色（十六进制格式）
        """
        fill = WordAPI._check_color(fill) if fill else None
        if not fill:
            return
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), fill)
        shading.set(qn("w:val"), "clear")

    @staticmethod
    def _set_cell_margins(
        cell: Any,
        top: int = 80,
        start: int = 80,
        bottom: int = 80,
        end: int = 80,
    ) -> None:
        """设置单元格内边距。

        :param cell: python-docx 的 Cell 对象
        :param top: 上边距（单位：twips）
        :param start: 左边距（单位：twips）
        :param bottom: 下边距（单位：twips）
        :param end: 右边距（单位：twips）
        """
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_margin = tc_pr.find(qn("w:tcMar"))
        if tc_margin is None:
            tc_margin = OxmlElement("w:tcMar")
            tc_pr.append(tc_margin)
        for key, value in {
            "top": top,
            "start": start,
            "bottom": bottom,
            "end": end,
        }.items():
            node = tc_margin.find(qn(f"w:{key}"))
            if node is None:
                node = OxmlElement(f"w:{key}")
                tc_margin.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    @staticmethod
    def _cm_to_dxa(cm_value: float) -> int:
        """将厘米转换为 DXA 单位（Word 内部单位）。

        :param cm_value: 厘米值
        :return: DXA 单位的整数值
        """
        return int(cm_value / 2.54 * 1440)

    @staticmethod
    def _set_cell_width(cell: Any, width_cm: float) -> None:
        """设置单元格宽度。

        :param cell: python-docx 的 Cell 对象
        :param width_cm: 宽度（厘米）
        """
        width_dxa = WordAPI._cm_to_dxa(width_cm)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_width = tc_pr.find(qn("w:tcW"))
        if tc_width is None:
            tc_width = OxmlElement("w:tcW")
            tc_pr.append(tc_width)
        tc_width.set(qn("w:w"), str(width_dxa))
        tc_width.set(qn("w:type"), "dxa")

    @staticmethod
    def _set_table_fixed_layout(table: Any) -> None:
        """设置表格为固定列宽布局。

        :param table: python-docx 的 Table 对象
        """
        table_pr = table._tbl.tblPr
        table_layout = table_pr.find(qn("w:tblLayout"))
        if table_layout is None:
            table_layout = OxmlElement("w:tblLayout")
            table_pr.append(table_layout)
        table_layout.set(qn("w:type"), "fixed")

    @staticmethod
    def _set_table_grid_widths(table: Any, col_widths_cm: Sequence[float]) -> None:
        """设置表格网格列宽。

        :param table: python-docx 的 Table 对象
        :param col_widths_cm: 各列宽度列表（厘米）
        """
        table_grid = table._tbl.tblGrid
        if table_grid is None:
            table_grid = OxmlElement("w:tblGrid")
            table._tbl.insert(1, table_grid)
        else:
            for child in list(table_grid):
                table_grid.remove(child)

        for width_cm in col_widths_cm:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(WordAPI._cm_to_dxa(float(width_cm))))
            table_grid.append(grid_col)

    @staticmethod
    def _clear_paragraph(paragraph: Any) -> None:
        paragraph_element = paragraph._element
        for child in list(paragraph_element):
            paragraph_element.remove(child)

    @classmethod
    def _is_image_file(cls, path_str: str) -> bool:
        return Path(path_str).suffix.lower() in cls.IMAGE_EXTENSIONS

    @staticmethod
    def _safe_image_size(path_str: str) -> tuple[int, int]:
        with Image.open(path_str) as image:
            return image.size

    @staticmethod
    def _path_to_str(value: Any) -> str:
        return str(value)

    def _is_existing_image_path(self, value: Any) -> bool:
        if not isinstance(value, (str, os.PathLike)):
            return False
        path_str = self._path_to_str(value)
        return os.path.exists(path_str) and self._is_image_file(path_str)

    @staticmethod
    def _image_path_from_part(part: Any) -> str | None:
        if not isinstance(part, dict):
            return None
        image_path = part.get("image") or part.get("path")
        if not image_path:
            return None
        if part.get("type") not in (None, "image"):
            return None
        return str(image_path)

    @staticmethod
    def _text_from_part(part: Any) -> Any:
        if not isinstance(part, dict):
            return part
        if part.get("type") not in (None, "text"):
            return None
        if "text" in part:
            return part.get("text")
        if "value" in part:
            return part.get("value")
        return None

    def _cell_parts_from_value(self, value: Any) -> list[Any] | None:
        if isinstance(value, list):
            return value
        if isinstance(value, dict) and value.get("type") == "mixed":
            parts = value.get("parts", [])
            if not isinstance(parts, list):
                raise TypeError("混排单元格的 parts 必须是列表")
            return parts
        return None

    def _add_picture_run(
        self,
        paragraph: Any,
        image_path: str,
        table_style: TableStyle,
        width_cm: float | None = None,
        height_cm: float | None = None,
    ) -> Any:
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"图片不存在: {image_path}")
        if not self._is_image_file(image_path):
            raise ValueError(f"不是支持的图片文件: {image_path}")

        run = paragraph.add_run()
        if width_cm is not None and height_cm is not None:
            run.add_picture(image_path, width=Cm(float(width_cm)), height=Cm(float(height_cm)))
        elif width_cm is not None:
            run.add_picture(image_path, width=Cm(float(width_cm)))
        elif height_cm is not None:
            run.add_picture(image_path, height=Cm(float(height_cm)))
        elif table_style.auto_fit_image:
            image_width_px, image_height_px = self._safe_image_size(image_path)
            effective_width_cm = table_style.image_width_cm
            effective_height_cm = (
                effective_width_cm * image_height_px / image_width_px
                if image_width_px
                else effective_width_cm
            )
            run.add_picture(
                image_path,
                width=Cm(effective_width_cm),
                height=Cm(effective_height_cm),
            )
        else:
            run.add_picture(image_path, width=Cm(table_style.image_width_cm))
        return run

    def _add_cell_text_part(
        self,
        paragraph: Any,
        value: Any,
        style: Optional[CellStyle],
    ) -> Any:
        text = "" if value is None else str(value)
        run = paragraph.add_run(text)
        if style is not None:
            self._set_run_font(
                run,
                font_name=style.font_name,
                font_size=style.font_size,
                bold=style.bold,
                italic=style.italic,
                font_color=style.font_color,
            )
        return run

    def _add_cell_part(
        self,
        paragraph: Any,
        part: Any,
        style: Optional[CellStyle],
        table_style: TableStyle,
    ) -> None:
        image_path = self._image_path_from_part(part)
        if image_path is not None:
            self._add_picture_run(
                paragraph,
                image_path,
                table_style,
                width_cm=part.get("width_cm"),
                height_cm=part.get("height_cm"),
            )
            return

        if self._is_existing_image_path(part):
            self._add_picture_run(paragraph, self._path_to_str(part), table_style)
            return

        nested_parts = self._cell_parts_from_value(part)
        if nested_parts is not None:
            for nested_part in nested_parts:
                self._add_cell_part(paragraph, nested_part, style, table_style)
            return

        text_value = self._text_from_part(part)
        if text_value is None and isinstance(part, dict):
            text_value = part
        self._add_cell_text_part(paragraph, text_value, style)

    def add_empty_paragraph(
        self,
        container: Any,
        style: Optional[TextStyle] = None,
    ) -> Any:
        paragraph = container.add_paragraph()
        self._apply_paragraph_style(paragraph, getattr(style, "style_name", None))
        self._apply_paragraph_direct_format(paragraph, style)
        return paragraph

    def add_text_run(
        self,
        paragraph: Any,
        text: TextValue,
        style: Optional[TextStyle] = None,
    ) -> Any:
        run = paragraph.add_run("" if text is None else str(text))
        if style is not None:
            self._set_run_font(
                run,
                font_name=style.font_name,
                font_size=style.font_size,
                bold=style.bold,
                italic=style.italic,
                font_color=style.font_color,
            )
        return run

    def add_paragraph(
        self,
        container: Any,
        text: TextValue = "",
        style: Optional[TextStyle] = None,
    ) -> Any:
        paragraph = self.add_empty_paragraph(container, style)
        self.add_text_run(paragraph, text, style)
        return paragraph

    def add_image_block(
        self,
        container: Any,
        image_path: str,
        width_cm: Optional[float] = None,
        height_cm: Optional[float] = None,
        align: str = "center",
        style: Optional[TextStyle] = None,
    ) -> Any:
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"图片不存在: {image_path}")

        paragraph = self.add_empty_paragraph(container, style or IMAGE_STYLE)
        paragraph.alignment = self._get_paragraph_alignment(align)
        run = paragraph.add_run()

        if width_cm and height_cm:
            run.add_picture(image_path, width=Cm(width_cm), height=Cm(height_cm))
        elif width_cm:
            run.add_picture(image_path, width=Cm(width_cm))
        elif height_cm:
            run.add_picture(image_path, height=Cm(height_cm))
        else:
            run.add_picture(image_path)

        return paragraph

    def _apply_cell_style(
        self,
        cell: Any,
        style: Optional[CellStyle],
    ) -> None:
        style = style or make_cell_style()
        cell.vertical_alignment = self._get_vertical_alignment(style.vertical_align)
        self._set_cell_margins(cell)

        if style.bg_color:
            self._set_cell_background(cell, style.bg_color)

        paragraphs = cell.paragraphs or [cell.add_paragraph()]
        for index, paragraph in enumerate(paragraphs):
            if index == 0:
                self._clear_paragraph(paragraph)
            self._apply_paragraph_style(paragraph, style.paragraph_style_name)
            if style.align is not None:
                paragraph.alignment = self._get_paragraph_alignment(style.align)
            if style.line_spacing is not None:
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing = style.line_spacing
                paragraph_format.line_spacing_rule = (
                    WD_LINE_SPACING.SINGLE
                    if style.line_spacing == 1
                    else WD_LINE_SPACING.MULTIPLE
                )

    def _fill_cell_value(
        self,
        cell: Any,
        value: Any,
        style: Optional[CellStyle],
        table_style: TableStyle,
    ) -> None:
        self._apply_cell_style(cell, style)
        paragraph = cell.paragraphs[0]

        if self._is_existing_image_path(value):
            self._add_picture_run(paragraph, self._path_to_str(value), table_style)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            return

        image_path = self._image_path_from_part(value)
        if image_path is not None:
            self._add_picture_run(
                paragraph,
                image_path,
                table_style,
                width_cm=value.get("width_cm"),
                height_cm=value.get("height_cm"),
            )
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            return

        parts = self._cell_parts_from_value(value)
        if parts is not None:
            for part in parts:
                self._add_cell_part(paragraph, part, style, table_style)
            return

        self._add_cell_text_part(paragraph, self._text_from_part(value), style)

    def _set_table_borders(
        self,
        table: Any,
        border_color: str = "000000",
        border_size: str = "8",
    ) -> None:
        table_properties = table._tbl.tblPr
        borders = table_properties.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            table_properties.append(borders)

        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            node = borders.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                borders.append(node)
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), str(border_size))
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), border_color)

    def insert_table(
        self,
        container: Any,
        data: list[list[TextValue]],
        header_style: Optional[CellStyle] = None,
        body_style: Optional[CellStyle] = None,
        table_style: Optional[TableStyle] = None,
    ) -> Any:
        if not data:
            raise ValueError("表格数据不能为空")

        col_count = max(len(row) for row in data)
        row_count = len(data)
        table_style = table_style or make_table_style()
        header_style = header_style or make_cell_style(
            paragraph_style_name="KL表格表头",
            align="center",
        )
        body_style = body_style or make_cell_style(
            paragraph_style_name="KL表格文字",
            align="left",
        )

        table = container.add_table(rows=row_count, cols=col_count)
        table.alignment = self._get_table_alignment(table_style.align)
        table.autofit = False
        self._set_table_fixed_layout(table)
        self._set_table_borders(table, table_style.border_color, table_style.border_size)

        if table_style.col_widths_cm:
            effective_widths = [
                float(width_cm)
                for width_cm in table_style.col_widths_cm[:col_count]
            ]
            if len(effective_widths) < col_count:
                effective_widths.extend(
                    [float(effective_widths[-1])] * (col_count - len(effective_widths))
                )

            self._set_table_grid_widths(table, effective_widths)

            for col_idx, width_cm in enumerate(effective_widths):
                for row in table.rows:
                    self._set_cell_width(row.cells[col_idx], width_cm)

        if table_style.row_heights_cm:
            for row_idx, height_cm in enumerate(table_style.row_heights_cm):
                if row_idx >= row_count or height_cm is None:
                    continue
                row = table.rows[row_idx]
                row.height = Cm(float(height_cm))
                row.height_rule = (
                    WD_ROW_HEIGHT_RULE.EXACTLY
                    if table_style.exact_row_height
                    else WD_ROW_HEIGHT_RULE.AT_LEAST
                )

        for row_index, row_data in enumerate(data):
            for col_index in range(col_count):
                value = row_data[col_index] if col_index < len(row_data) else ""
                style = (
                    header_style
                    if row_index < table_style.header_rows
                    else body_style
                )
                self._fill_cell_value(
                    table.cell(row_index, col_index),
                    value,
                    style,
                    table_style,
                )

        return table

    def insert_table_by_config(self, container: Any, table_config: TableConfig) -> Any:
        default_data = [
            ["序号", "内容"],
            ["未检测到数据填充", "未检测到数据填充"],
        ]

        max_table_width_cm = 14.0
        default_row_height_cm = 0.6
        default_col_width_cm = 4.0
        min_col_width_cm = 1.2

        data = table_config.get("data") or default_data
        data = self._normalize_table_data(data) or default_data

        row_count = len(data)
        input_row_heights = table_config.get("row_heights_cm") or []
        input_col_widths = table_config.get("col_widths_cm") or []

        row_heights_cm = [
            input_row_heights[i] if i < len(input_row_heights) else default_row_height_cm
            for i in range(row_count)
        ]

        col_widths_cm = self._build_col_widths(
            data=data,
            input_col_widths=input_col_widths,
            max_table_width_cm=max_table_width_cm,
            default_col_width_cm=default_col_width_cm,
            min_col_width_cm=min_col_width_cm,
        )

        style_config = table_config.get("style") or {}

        header_style = style_config.get("header", TABLE_HEADER_STYLE)
        body_style = style_config.get("body", TABLE_BODY_STYLE)
        table_style = style_config.get(
            "table",
            make_table_style(
                col_widths_cm=col_widths_cm,
                row_heights_cm=row_heights_cm,
            ),
        )

        return self.insert_table(
            container=container,
            data=data,
            header_style=header_style,
            body_style=body_style,
            table_style=table_style,
        )

    def _normalize_table_data(self, data: list[list[Any]]) -> list[list[Any]]:
        if not data:
            return []

        col_count = max(len(row) for row in data)
        normalized = []
        for row in data:
            new_row = ["" if cell is None else cell for cell in row]
            if len(new_row) < col_count:
                new_row.extend([""] * (col_count - len(new_row)))
            normalized.append(new_row)
        return normalized

    def _cell_display_text(self, value: Any) -> str:
        image_path = self._image_path_from_part(value)
        if image_path is not None or self._is_existing_image_path(value):
            return ""

        parts = self._cell_parts_from_value(value)
        if parts is not None:
            return "".join(self._cell_display_text(part) for part in parts)

        text_value = self._text_from_part(value)
        if text_value is None and isinstance(value, dict):
            text_value = value
        return "" if text_value is None else str(text_value)

    def _build_col_widths(
        self,
        data: list[list[Any]],
        input_col_widths: list[float],
        max_table_width_cm: float = 14.0,
        default_col_width_cm: float = 4.0,
        min_col_width_cm: float = 1.2,
    ) -> list[float]:
        if not data or not data[0]:
            return []

        col_count = len(data[0])

        col_weights = []
        for col_idx in range(col_count):
            max_len = 1
            for row in data:
                cell_text = self._cell_display_text(row[col_idx] if col_idx < len(row) else "")
                text_len = self._get_display_length(cell_text)
                max_len = max(max_len, text_len)
            col_weights.append(max_len)

        total_weight = sum(col_weights) or col_count
        auto_widths_all = [
            max(min_col_width_cm, max_table_width_cm * weight / total_weight)
            for weight in col_weights
        ]
        auto_widths_all = self._scale_widths_to_max(auto_widths_all, max_table_width_cm)

        specified: list[Optional[float]] = []
        unspecified_indexes = []
        for i in range(col_count):
            if i < len(input_col_widths) and input_col_widths[i] is not None and input_col_widths[i] > 0:
                specified.append(float(input_col_widths[i]))
            else:
                specified.append(None)
                unspecified_indexes.append(i)

        if len(unspecified_indexes) == col_count:
            return auto_widths_all

        if not unspecified_indexes:
            final_widths = [w if w is not None else default_col_width_cm for w in specified]
            return self._scale_widths_to_max(final_widths, max_table_width_cm)

        specified_sum = sum(w for w in specified if w is not None)

        if specified_sum >= max_table_width_cm:
            desired = [
                specified[i] if specified[i] is not None else auto_widths_all[i]
                for i in range(col_count)
            ]
            return self._scale_widths_to_max(desired, max_table_width_cm)

        remaining_width = max_table_width_cm - specified_sum
        unspecified_weights = [col_weights[i] for i in unspecified_indexes]
        unspecified_total_weight = sum(unspecified_weights) or len(unspecified_indexes)

        final_widths = []
        for i in range(col_count):
            if specified[i] is not None:
                final_widths.append(specified[i])  # type: ignore[arg-type]
            else:
                weight = col_weights[i]
                width = remaining_width * weight / unspecified_total_weight
                final_widths.append(max(min_col_width_cm, width))

        return self._scale_widths_to_max(final_widths, max_table_width_cm)

    def _scale_widths_to_max(self, widths: list[float], max_total: float) -> list[float]:
        if not widths:
            return widths

        total = sum(widths)
        if total <= 0 or total <= max_total:
            return widths

        ratio = max_total / total
        return [round(w * ratio, 4) for w in widths]

    def _get_display_length(self, text: Any) -> int:
        s = "" if text is None else str(text)
        length = 0
        for ch in s:
            length += 2 if unicodedata.east_asian_width(ch) in ("F", "W", "A") else 1
        return max(length, 1)

    def add_field_run(self, paragraph: Any, field_code: str) -> Any:
        run_begin = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run_begin._r.append(fld_char_begin)

        run_instr = paragraph.add_run()
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = field_code
        run_instr._r.append(instr_text)

        run_sep = paragraph.add_run()
        fld_char_sep = OxmlElement("w:fldChar")
        fld_char_sep.set(qn("w:fldCharType"), "separate")
        run_sep._r.append(fld_char_sep)

        placeholder_run = paragraph.add_run("")

        run_end = paragraph.add_run()
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run_end._r.append(fld_char_end)
        return placeholder_run

    def add_field_paragraph(
        self,
        container: Any,
        parts: list[PartConfig],
        style: Optional[TextStyle] = None,
    ) -> Any:
        paragraph = self.add_empty_paragraph(container, style or BODY_STYLE)
        run_style = style or BODY_STYLE

        for part in parts:
            part_type = part.get("type", "text")
            if part_type == "text":
                self.add_text_run(paragraph, part.get("value", ""), run_style)
            elif part_type == "field":
                self.add_field_run(paragraph, part.get("code", ""))
            else:
                raise ValueError(f"不支持的 part type: {part_type}")

        return paragraph

    def add_page_footer(
        self,
        container: Any,
        style: Optional[TextStyle] = None,
    ) -> Any:
        return self.add_field_paragraph(
            container,
            [
                {"type": "text", "value": "第 "},
                {"type": "field", "code": "PAGE"},
                {"type": "text", "value": " 页 / 共 "},
                {"type": "field", "code": "NUMPAGES"},
                {"type": "text", "value": " 页"},
            ],
            style or FOOTER_STYLE,
        )

    def _build_auto_caption_parts(
        self,
        label: str,
        seq_name: str,
        title: str,
    ) -> list[PartConfig]:
        parts: list[PartConfig] = [
            {"type": "text", "value": f"{label} "},
            {"type": "field", "code": r"STYLEREF KL一级标题 \n \* MERGEFORMAT "},
            {"type": "text", "value": "-"},
            {"type": "field", "code": rf"SEQ {seq_name} \* ARABIC \s 1"},
        ]
        if title:
            parts.append({"type": "text", "value": f" {title}"})
        return parts

    def add_figure_caption_auto(
        self,
        container: Any,
        title: str,
        style: Optional[TextStyle] = None,
    ) -> Any:
        return self.add_field_paragraph(
            container,
            self._build_auto_caption_parts("图", "图", title),
            style or CAPTION_STYLE,
        )

    def add_table_caption_auto(
        self,
        container: Any,
        title: str,
        style: Optional[TextStyle] = None,
    ) -> Any:
        return self.add_field_paragraph(
            container,
            self._build_auto_caption_parts("表", "表", title),
            style or CAPTION_STYLE,
        )

    def render(self, context: dict[str, Any], output_path: str) -> str:
        output_path = str(output_path)
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        self.doc.render(context)
        self.doc.save(output_path)

        return output_path

    def _write_single_header(
        self,
        section: Any,
        text: str,
        style: Optional[TextStyle],
    ) -> None:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        self._clear_paragraph(paragraph)
        self._apply_paragraph_style(paragraph, getattr(style, "style_name", None))
        self._apply_paragraph_direct_format(paragraph, style)
        paragraph.alignment = self._get_paragraph_alignment(
            getattr(style, "align", "center")
        )
        self.add_text_run(paragraph, text, style)

    def _set_section_page_start(
        self,
        section: Any,
        start: int = 1,
    ) -> None:
        sect_pr = section._sectPr
        old_pg_num_type = sect_pr.find(qn("w:pgNumType"))
        if old_pg_num_type is not None:
            sect_pr.remove(old_pg_num_type)

        pg_num_type = OxmlElement("w:pgNumType")
        pg_num_type.set(qn("w:start"), str(start))
        sect_pr.append(pg_num_type)

    def _remove_section_page_start(self, section: Any) -> None:
        sect_pr = section._sectPr
        old_pg_num_type = sect_pr.find(qn("w:pgNumType"))
        if old_pg_num_type is not None:
            sect_pr.remove(old_pg_num_type)

    def _write_single_footer(
        self,
        section: Any,
        style: Optional[TextStyle],
    ) -> None:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        self._clear_footer_paragraph_keep_ppr(paragraph)

        self._apply_paragraph_style(paragraph, getattr(style, "style_name", None))
        self._apply_paragraph_direct_format(paragraph, style)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run1 = paragraph.add_run("第")
        if style is not None:
            self._set_run_font(
                run1,
                font_name=style.font_name,
                font_size=style.font_size,
                bold=style.bold,
                italic=style.italic,
                font_color=style.font_color,
            )

        self._add_footer_field_run_demo(paragraph, "PAGE")

        run2 = paragraph.add_run("页 共")
        if style is not None:
            self._set_run_font(
                run2,
                font_name=style.font_name,
                font_size=style.font_size,
                bold=style.bold,
                italic=style.italic,
                font_color=style.font_color,
            )

        self._add_footer_field_run_demo(paragraph, "SECTIONPAGES")

        run3 = paragraph.add_run("页")
        if style is not None:
            self._set_run_font(
                run3,
                font_name=style.font_name,
                font_size=style.font_size,
                bold=style.bold,
                italic=style.italic,
                font_color=style.font_color,
            )

    def write_header_footer(
            self,
            docx_path: str,
            header_text: Optional[str] = None,
            header_style: Optional[TextStyle] = None,
            footer_style: Optional[TextStyle] = None,
    ) -> str:
        document = Document(docx_path)

        if not document.sections:
            document.save(docx_path)
            return docx_path

        # 1. 页眉：全局所有节统一写
        if header_text is not None:
            for section in document.sections:
                section.header.is_linked_to_previous = False
                self._write_single_header(
                    section,
                    header_text,
                    header_style or HEADER_STYLE,
                )

        # 2. 页脚：只处理最后一节
        last_section = document.sections[-1]
        last_section.footer.is_linked_to_previous = False
        self._set_section_page_start(last_section, start=1)
        self._write_single_footer(
            last_section,
            footer_style or FOOTER_STYLE,
        )

        document.save(docx_path)
        return docx_path

    def _add_footer_field_run_demo(self, paragraph: Any, field_name: str) -> None:
        run = paragraph.add_run()

        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = field_name

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)

    def _clear_footer_paragraph_keep_ppr(self, paragraph: Any) -> None:
        p = paragraph._p
        for child in list(p):
            if child.tag != qn("w:pPr"):
                p.remove(child)

# if __name__ == '__main__':
#     pass
__all__ = [
    "TextValue",
    "PartConfig",
    "TableConfig",
    "DocContainer",
    "WordAPI",
]
