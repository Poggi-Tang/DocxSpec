"""KL 标准 Word 文档处理能力。

本模块把公司 Word 标准化流程沉淀为 DocxSpec 的库能力。它不绑定 GUI，
只提供可复用的 DOCX 分类、样式规范化、题注字段、标准文档构建和检查入口。
"""

from __future__ import annotations

import json
import re
import shutil
import tempfile
import zipfile
from collections import defaultdict, deque
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docxcompose.composer import Composer

PACKAGE_ASSETS_DIR = Path(__file__).resolve().parent / "assets"
DEFAULT_MASTER_TEMPLATE = PACKAGE_ASSETS_DIR / "docxspec_master_template.docx"

STYLE_MAIN_TITLE = "KL主标题"
STYLE_HEADING_1 = "KL一级标题"
STYLE_HEADING_2 = "KL二级标题"
STYLE_HEADING_3 = "KL三级标题"
STYLE_HEADING_4 = "KL四级标题"
STYLE_HEADING_5 = "KL五级标题"
STYLE_HEADING_6 = "KL六级标题"
STYLE_BODY = "KL正文"
STYLE_CAPTION = "KL题注"
STYLE_TABLE_HEADER = "KL表格表头"
STYLE_TABLE_BODY = "KL表格文字"
STYLE_IMAGE = "KL图片"

REQUIRED_STYLES = [
    STYLE_HEADING_1,
    STYLE_HEADING_2,
    STYLE_HEADING_3,
    STYLE_HEADING_4,
    STYLE_HEADING_5,
    STYLE_HEADING_6,
    STYLE_BODY,
    STYLE_CAPTION,
    STYLE_TABLE_HEADER,
    STYLE_TABLE_BODY,
]

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "v": "urn:schemas-microsoft-com:vml",
}

STYLE_TO_TYPE = {
    STYLE_HEADING_1: ("heading_1", "high"),
    "Heading 1": ("heading_1", "high"),
    "标题 1": ("heading_1", "high"),
    "1": ("heading_1", "medium"),
    STYLE_HEADING_2: ("heading_2", "high"),
    "Heading 2": ("heading_2", "high"),
    "标题 2": ("heading_2", "high"),
    "2": ("heading_2", "medium"),
    STYLE_HEADING_3: ("heading_3", "high"),
    "Heading 3": ("heading_3", "high"),
    "标题 3": ("heading_3", "high"),
    "3": ("heading_3", "medium"),
    STYLE_HEADING_4: ("heading_4", "high"),
    "Heading 4": ("heading_4", "high"),
    "标题 4": ("heading_4", "high"),
    "4": ("heading_4", "medium"),
    STYLE_HEADING_5: ("heading_5", "high"),
    "Heading 5": ("heading_5", "high"),
    "标题 5": ("heading_5", "high"),
    "5": ("heading_5", "medium"),
    STYLE_HEADING_6: ("heading_6", "high"),
    "Heading 6": ("heading_6", "high"),
    "标题 6": ("heading_6", "high"),
    "6": ("heading_6", "medium"),
}

CHINESE_NUM = "一二三四五六七八九十百千万〇零两"
H1_EXACT_TITLES = {"需求背景", "从需求分析到方案落地", "示例演示", "总结"}
FRONT_MATTER_TITLES = {
    "目录",
    "目 录",
    "修订记录",
    "版本记录",
    "变更记录",
    "图目录",
    "图录",
    "插图清单",
    "表目录",
    "表录",
    "表格清单",
}

NUMBER_PREFIX_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("heading_1", re.compile(rf"^第?\s*[{CHINESE_NUM}]+\s*[章节篇部分]\s*[、.．:：]?\s*")),
    ("heading_1", re.compile(rf"^[{CHINESE_NUM}]+[、.．]\s*")),
    ("heading_1", re.compile(r"^\d+[、．]\s*|^\d+\s+")),
    ("heading_6", re.compile(r"^\d+[.．]\d+[.．]\d+[.．]\d+[.．]\d+[.．]\d+[、.．]?\s*")),
    ("heading_5", re.compile(r"^\d+[.．]\d+[.．]\d+[.．]\d+[.．]\d+[、.．]?\s*")),
    ("heading_4", re.compile(r"^\d+[.．]\d+[.．]\d+[.．]\d+[、.．]?\s*")),
    ("heading_3", re.compile(r"^\d+[.．]\d+[.．]\d+[、.．]?\s*")),
    ("heading_2", re.compile(r"^\d+[.．]\d+[、.．]?\s*")),
]

CAPTION_TEXT_PATTERNS = {
    "figure_caption": re.compile(r"^\s*图\s*[\d一二三四五六七八九十]+[-－—.．]\d+\s+\S+"),
    "table_caption": re.compile(r"^\s*表\s*[\d一二三四五六七八九十]+[-－—.．]\d+\s+\S+"),
}

LIST_ITEM_PREFIX_PATTERN = re.compile(
    rf"^((?:[（(]\s*(?:\d+|[{CHINESE_NUM}]+)\s*[）)])|(?:\d+|[{CHINESE_NUM}]+)[）)])\s*"
)
LIST_ITEM_BLOCK_KEYWORDS = {
    "端口",
    "选项",
    "步骤",
    "条件",
    "参数值",
    "输入信号",
    "输出信号",
    "信号",
    "取值",
    "默认值",
}


@dataclass
class BlockReport:
    """DOCX 正文块分类结果。"""

    index: int
    kind: str
    type: str
    text: str = ""
    style_id: str | None = None
    style_name: str | None = None
    reason: str | None = None
    confidence: str | None = None
    has_figure: bool = False
    field_codes: list[str] | None = None
    rows: int | None = None
    columns: int | None = None
    caption: dict[str, Any] | None = None


def iter_block_items(doc: Document) -> list[Paragraph | Table]:
    """按 Word 正文顺序返回段落和表格块。"""
    blocks: list[Paragraph | Table] = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            blocks.append(Paragraph(child, doc))
        elif child.tag == qn("w:tbl"):
            blocks.append(Table(child, doc))
    return blocks


def paragraph_has_image(paragraph: Paragraph) -> bool:
    """判断段落中是否包含图片。"""
    return bool(paragraph._element.xpath(".//w:drawing") or paragraph._element.xpath(".//w:pict"))


def paragraph_field_xml(paragraph: Paragraph) -> str:
    """返回段落 XML，用于检查域代码。"""
    return paragraph._element.xml


def style_name(paragraph: Paragraph) -> str:
    """安全获取段落样式名称。"""
    try:
        return paragraph.style.name or ""
    except Exception:
        return ""


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """替换段落纯文本。"""
    paragraph.text = text


def clear_numbering(paragraph: Paragraph) -> None:
    """清除段落自身编号设置，让 KL 标题样式接管编号。"""
    ppr = paragraph._element.get_or_add_pPr()
    for num_pr in list(ppr.findall(qn("w:numPr"))):
        ppr.remove(num_pr)


def normalize_paragraph_to_style(paragraph: Paragraph, style: str) -> None:
    """应用段落样式，并清理直接格式和非样式段落属性。"""
    paragraph.style = style
    ppr = paragraph._element.get_or_add_pPr()
    for child in list(ppr):
        if child.tag != qn("w:pStyle"):
            ppr.remove(child)
    for run in paragraph.runs:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is not None:
            run._element.remove(rpr)


def normalize_table_cell_paragraph(paragraph: Paragraph, style: str) -> None:
    """应用表格单元格段落样式。"""
    normalize_paragraph_to_style(paragraph, style)


def clean_heading_number(text: str) -> tuple[str, str | None]:
    """移除常见手工标题序号，并返回推断的标题类型。"""
    stripped = text.strip()
    for kind, pattern in NUMBER_PREFIX_PATTERNS:
        cleaned = pattern.sub("", stripped, count=1).strip()
        if cleaned != stripped and cleaned:
            return cleaned, kind
    return stripped, None


def normalize_heading_text(text: str) -> str:
    """清理标题文本中的制表符、换行符和多余空白。"""
    return re.sub(r"\s+", " ", text).strip()


def is_code_like_line(text: str) -> bool:
    """判断段落是否更像伪代码/代码行，而不是标题。"""
    raw = text.replace("\u00a0", " ").strip()
    if not raw:
        return False
    if re.fullmatch(r"[{}()\[\];,]+", raw):
        return True
    if raw.startswith(("}", "{", "//", "/*", "*/", "*", "#", "||", "&&", "|", "&")):
        return True
    if re.match(r"^(<=|>=|==|!=|<|>|\+|-|\*|/)", raw):
        return True
    if re.match(r"^(if|else|for|while|switch|case|return|break|continue)\b", raw):
        return True
    if re.match(r"^[_A-Za-z][\w.\->\[\]]*\s*(=|\+=|-=|\*=|/=|==|!=|<=|>=|<|>)", raw):
        return True
    if re.match(r"^[_A-Za-z][\w.\->\[\]]*\s*\(.*\)\s*;?$", raw):
        return True
    if re.search(r"[_A-Za-z][\w.\->\[\]]*\s*=", raw):
        return True
    if re.search(r"->|\+\+|--", raw):
        return True
    if raw.endswith((";", "{", "}")) and re.search(r"[_A-Za-z][\w.\->\[\]]*", raw):
        return True
    return False


def split_list_item_prefix(text: str) -> tuple[str, str] | None:
    """拆分 1）/（1）/一）这类单层编号列表项。"""
    match = LIST_ITEM_PREFIX_PATTERN.match(text.strip())
    if not match:
        return None
    return match.group(1), text.strip()[match.end() :].strip()


def is_probable_list_item(text: str) -> bool:
    """判断单层编号行是否更像正文列表项，而不是章节标题。"""
    parsed = split_list_item_prefix(text)
    if not parsed:
        return False
    _prefix, body = parsed
    if not body:
        return True
    if len(body) > 18:
        return True
    return any(keyword in body for keyword in LIST_ITEM_BLOCK_KEYWORDS)


def is_short_heading_candidate(text: str) -> bool:
    """判断短文本是否可作为普通标题候选。"""
    if not text or len(text) > 32:
        return False
    if "http://" in text or "https://" in text:
        return False
    if "：" in text or ":" in text:
        return False
    if text.endswith(("。", "；", ";", "，", ",")):
        return False
    return True


def classify_heading(text: str, paragraph: Paragraph | None = None) -> tuple[str, str, str, str]:
    """按文本和已有样式判断段落应使用的 KL 标题/正文类型。"""
    raw = text.strip()
    if is_code_like_line(raw):
        return "body", raw, "code-like line is body, not heading", "high"
    if re.match(r"^[①②③④⑤⑥⑦⑧⑨⑩]", raw):
        return "body", raw, "circled-number item is body, not heading", "high"
    if is_probable_list_item(raw):
        return "body", raw, "single-level numbered item is body pending outline analysis", "high"
    cleaned, numbered_kind = clean_heading_number(raw)
    if numbered_kind:
        if len(cleaned) > 42:
            return "body", raw, "numbered text too long to classify as heading", "medium"
        return numbered_kind, cleaned, "matched numbered heading prefix", "high"

    if paragraph is not None:
        raw_style = style_name(paragraph)
        normalized = raw_style.replace(" ", "").lower()
        if normalized == "listparagraph":
            return "body", raw, "list paragraph is body, not standalone heading", "high"
        style_map = {
            STYLE_HEADING_1: "heading_1",
            STYLE_HEADING_2: "heading_2",
            STYLE_HEADING_3: "heading_3",
            STYLE_HEADING_4: "heading_4",
            STYLE_HEADING_5: "heading_5",
            STYLE_HEADING_6: "heading_6",
        }
        if raw_style in style_map:
            return style_map[raw_style], raw, f"matched {raw_style} style", "high"
        for level in range(1, 7):
            if normalized in {f"heading{level}", f"标题{level}", str(level)}:
                return f"heading_{level}", raw, f"matched heading-{level} style", "high"

    if raw in H1_EXACT_TITLES:
        return "heading_1", raw, "matched known major section title", "high"
    return "body", raw, "default body paragraph", "low"


def document_has_figures(docx: Path) -> bool:
    """判断文档是否包含图片。"""
    doc = Document(str(docx))
    return any(paragraph_has_image(p) for p in doc.paragraphs)


def document_has_tables(docx: Path) -> bool:
    """判断文档是否包含表格。"""
    doc = Document(str(docx))
    return bool(doc.tables)


def first_nonblank_paragraph_text(docx: Path) -> str | None:
    """读取文档第一个非空段落文本，通常作为封面标题候选。"""
    doc = Document(str(docx))
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            return clean_heading_number(text)[0]
    return None


def remove_block(block: Paragraph | Table) -> None:
    """从文档 XML 中删除段落或表格块。"""
    element = block._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_template_sample_body(doc: Document) -> int:
    """删除母版中的示例正文，保留封面、目录、图表清单等前置内容。"""
    blocks = iter_block_items(doc)
    start_idx: int | None = None
    for idx, block in enumerate(blocks):
        if isinstance(block, Paragraph):
            if block.text.strip() == STYLE_HEADING_1 or style_name(block) == STYLE_HEADING_1:
                start_idx = idx
                break
    if start_idx is None:
        return len(blocks)
    for block in blocks[start_idx:]:
        remove_block(block)
    return start_idx


def _is_front_matter_title(text: str) -> bool:
    """判断文本是否是封面之后、正文之前的目录或记录类标题。"""
    normalized = re.sub(r"\s+", "", text)
    return text.strip() in FRONT_MATTER_TITLES or normalized in {
        re.sub(r"\s+", "", title) for title in FRONT_MATTER_TITLES
    }


def _is_toc_like_block(block: dict[str, Any]) -> bool:
    """判断块是否包含目录、图录或表录域。"""
    fields = "\n".join(block.get("field_codes") or [])
    text = str(block.get("text", "")).strip()
    return "TOC" in fields or _is_front_matter_title(text)


def detect_source_front_matter(source_report: dict[str, Any]) -> dict[str, Any]:
    """识别源文档正文开始位置，用于跳过旧封面、目录、图录、表录等前置页。"""
    blocks = source_report.get("blocks", [])
    body_start = 0
    reason = "no front matter detected"
    for idx, block in enumerate(blocks):
        text = str(block.get("text", "")).strip()
        block_type = str(block.get("type", ""))
        if _is_toc_like_block(block):
            continue
        if block_type == "heading_1" and not _is_front_matter_title(text):
            body_start = idx
            reason = "first non-front-matter level-1 heading"
            break

    skipped = []
    for block in blocks[:body_start]:
        text = str(block.get("text", "")).strip()
        if _is_front_matter_title(text):
            skip_reason = "front matter title"
        elif _is_toc_like_block(block):
            skip_reason = "toc/list field or title"
        elif block.get("kind") == "table":
            skip_reason = "front matter table before body"
        elif block.get("type") == "figure":
            skip_reason = "front matter figure before body"
        elif text:
            skip_reason = "cover/front matter text before body"
        else:
            skip_reason = "front matter blank before body"
        skipped.append(
            {
                "index": block.get("index"),
                "kind": block.get("kind"),
                "type": block.get("type"),
                "text": text,
                "reason": skip_reason,
            }
        )

    return {
        "detected": body_start > 0,
        "body_start_block": body_start,
        "body_start_reason": reason,
        "skipped_blocks_count": len(skipped),
        "skipped_blocks": skipped,
    }


def infer_source_title(input_docx: Path, source_report: dict[str, Any], front_matter: dict[str, Any]) -> str | None:
    """从源文档推断封面标题，优先使用旧封面中连续的标题行。"""
    if not front_matter.get("detected"):
        blocks = source_report.get("blocks", [])
        if (
            blocks
            and str(blocks[0].get("type", "")).startswith("heading_")
            and (
                str(blocks[0].get("confidence", "")) == "high"
                or "numbered heading" in str(blocks[0].get("reason", ""))
                or "matched style" in str(blocks[0].get("reason", ""))
            )
        ):
            return input_docx.stem
        return first_nonblank_paragraph_text(input_docx)

    skipped = front_matter.get("skipped_blocks", [])
    best: list[str] = []
    current: list[str] = []
    for block in skipped:
        text = str(block.get("text", "")).strip()
        block_type = str(block.get("type", ""))
        if (
            not text
            or block.get("kind") != "paragraph"
            or block_type not in {"heading_1", "heading_2", "heading_3", "body"}
            or _is_front_matter_title(text)
            or "文件编号" in text
            or "公司" in text
            or re.search(r"\d{4}\s*年|\d{4}[-/.]\d{1,2}[-/.]\d{1,2}", text)
        ):
            if len("".join(current)) > len("".join(best)):
                best = current
            current = []
            continue
        current.append(text)
    if len("".join(current)) > len("".join(best)):
        best = current
    if best:
        return "".join(best)
    return first_nonblank_paragraph_text(input_docx)


def infer_source_file_number(source_report: dict[str, Any]) -> str | None:
    """从源文档前部提取文件编号行。"""
    for block in source_report.get("blocks", [])[:40]:
        text = str(block.get("text", "")).strip()
        if re.match(r"^文件编号\s*[:：]", text):
            return text
    return None


def write_source_body_only_docx(input_docx: Path, output_docx: Path, body_start: int) -> None:
    """复制源文档并删除正文开始前的前置块。"""
    shutil.copy2(input_docx, output_docx)
    if body_start <= 0:
        return
    doc = Document(str(output_docx))
    blocks = iter_block_items(doc)
    for block in blocks[:body_start]:
        remove_block(block)
    doc.save(str(output_docx))


def replace_cover_title(doc: Document, title: str | None) -> None:
    """把母版封面标题替换为源文档标题。"""
    if not title:
        return
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "XXXX系统":
            paragraph.text = title
            normalize_paragraph_to_style(paragraph, STYLE_MAIN_TITLE)
        elif text == "技术方案":
            paragraph.text = ""


def replace_cover_file_number(doc: Document, file_number: str | None) -> None:
    """把母版封面文件编号占位替换为源文档文件编号。"""
    if not file_number:
        return
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("文件编号：") or text.startswith("文件编号:"):
            paragraph.text = file_number
            normalize_paragraph_to_style(paragraph, STYLE_HEADING_5)
            return


def remove_template_front_matter_samples(doc: Document) -> None:
    """删除母版目录、图录、表录中的示例缓存项，保留域本身等待 Word/WPS 更新。"""
    sample_pattern = re.compile(r"^\s*\d+[-－—.．]\d+\s+(图注|表注)\d+\s*(\t+\d+)?\s*$")
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if not sample_pattern.match(text):
            continue
        if "TOC" in paragraph_field_xml(paragraph):
            for text_node in paragraph._element.xpath(".//w:t"):
                text_node.text = ""
            for tab_node in list(paragraph._element.xpath(".//w:tab")):
                parent = tab_node.getparent()
                if parent is not None:
                    parent.remove(tab_node)
        else:
            remove_block(paragraph)


def set_default_headers(doc: Document, title: str | None) -> None:
    """把所有节的页眉设置为标题。"""
    if not title:
        return
    for section in doc.sections:
        for header in (section.header, section.even_page_header):
            if not header.paragraphs:
                header.add_paragraph(title)
            for paragraph in header.paragraphs:
                paragraph.text = title


def insert_paragraph_after(paragraph: Paragraph, style: str | None = None) -> Paragraph:
    """在指定段落后插入新段落。"""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    return new_paragraph


def append_field(run: Any, instr: str, placeholder: str) -> None:
    """向 run 中追加 Word 域代码。"""
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r.append(begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    r.append(instr_text)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    r.append(separate)

    text = OxmlElement("w:t")
    text.text = placeholder
    r.append(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r.append(end)


def has_figure_caption_field(paragraph: Paragraph) -> bool:
    """判断段落是否包含图题注 SEQ 域。"""
    return "SEQ 图" in paragraph_field_xml(paragraph)


def has_table_caption_field(paragraph: Paragraph) -> bool:
    """判断段落是否包含表题注 SEQ 域。"""
    return "SEQ 表" in paragraph_field_xml(paragraph)


def has_standard_caption_fields(paragraph: Paragraph, seq_name: str) -> bool:
    """判断题注是否同时包含公司标准章节域和指定 SEQ 域。"""
    fields = paragraph_field_xml(paragraph)
    return "STYLEREF KL一级标题" in fields and f"SEQ {seq_name}" in fields


def extract_caption_title(paragraph: Paragraph, label: str) -> str:
    """从已有题注文本中提取题名，去掉图/表编号部分。"""
    text = paragraph.text.strip()
    text = re.sub(r"\s+", " ", text)
    # 兼容“表 5-1 标题”“图 3.2 标题”“表 1 标题”等缓存文本。
    title = re.sub(
        rf"^\s*{label}\s*[\d一二三四五六七八九十]+(?:\s*[-－—.．]\s*\d+)?\s*",
        "",
        text,
    ).strip()
    if title == text and title.startswith(label):
        title = title[len(label):].strip(" -－—.．\t")
    return title


def is_plain_figure_caption_text(text: str) -> bool:
    """判断图片后方短文本是否像未编号图题注。"""
    stripped = text.strip()
    if not stripped or len(stripped) > 40:
        return False
    if not stripped.endswith("图"):
        return False
    if stripped.startswith(("如", "上", "下", "本")):
        return False
    if any(mark in stripped for mark in "。；;，,：:"):
        return False
    return True


def is_table_caption_paragraph(paragraph: Paragraph) -> bool:
    """判断段落是否是表题注字段或可识别的表题注文本。"""
    return has_table_caption_field(paragraph) or bool(
        CAPTION_TEXT_PATTERNS["table_caption"].match(paragraph.text.strip())
    )


def is_figure_caption_paragraph(paragraph: Paragraph) -> bool:
    """判断段落是否是图题注字段或可识别的图题注文本。"""
    return has_figure_caption_field(paragraph) or bool(
        CAPTION_TEXT_PATTERNS["figure_caption"].match(paragraph.text.strip())
    )


def add_figure_caption_after(
    paragraph: Paragraph,
    caption_text: str,
    chapter_no: int = 1,
    figure_no: int = 1,
) -> None:
    """在图片段落后插入 KL 图题注和必要空行。"""
    caption = insert_paragraph_after(paragraph, STYLE_CAPTION)
    normalize_paragraph_to_style(caption, STYLE_CAPTION)
    caption.add_run("图 ")
    append_field(caption.add_run(), r" STYLEREF KL一级标题 \n \* MERGEFORMAT ", str(chapter_no))
    caption.add_run("-")
    append_field(caption.add_run(), r" SEQ 图 \* ARABIC \s 1 ", str(figure_no))
    caption.add_run(f" {caption_text}")
    blank = insert_paragraph_after(caption, STYLE_BODY)
    normalize_paragraph_to_style(blank, STYLE_BODY)


def set_figure_caption_paragraph(
    paragraph: Paragraph,
    caption_text: str,
    chapter_no: int = 1,
    figure_no: int = 1,
) -> None:
    """把图片后的未编号题注文本改为 KL 图题注字段。"""
    paragraph.text = ""
    normalize_paragraph_to_style(paragraph, STYLE_CAPTION)
    paragraph.add_run("图 ")
    append_field(paragraph.add_run(), r" STYLEREF KL一级标题 \n \* MERGEFORMAT ", str(chapter_no))
    paragraph.add_run("-")
    append_field(paragraph.add_run(), r" SEQ 图 \* ARABIC \s 1 ", str(figure_no))
    paragraph.add_run(f" {caption_text}")


def set_table_caption_paragraph(
    paragraph: Paragraph,
    caption_text: str,
    chapter_no: int = 1,
    table_no: int = 1,
) -> None:
    """把已有表题注文本或非标准字段改为 KL 标准表题注字段。"""
    paragraph.text = ""
    normalize_paragraph_to_style(paragraph, STYLE_CAPTION)
    paragraph.add_run("表 ")
    append_field(paragraph.add_run(), r" STYLEREF KL一级标题 \n \* MERGEFORMAT ", str(chapter_no))
    paragraph.add_run("-")
    append_field(paragraph.add_run(), r" SEQ 表 \* ARABIC \s 1 ", str(table_no))
    paragraph.add_run(f" {caption_text}")


def insert_paragraph_before(block: Paragraph | Table, style: str | None = None) -> Paragraph:
    """在指定块前插入新段落。"""
    new_p = OxmlElement("w:p")
    block._element.addprevious(new_p)
    paragraph = Paragraph(new_p, block._parent)
    if style:
        paragraph.style = style
    return paragraph


def add_table_caption_before(
    table: Table,
    caption_text: str,
    chapter_no: int = 1,
    table_no: int = 1,
) -> None:
    """在表格前插入 KL 表题注和必要空行。"""
    caption = insert_paragraph_before(table, STYLE_CAPTION)
    normalize_paragraph_to_style(caption, STYLE_CAPTION)
    caption.add_run("表 ")
    append_field(caption.add_run(), r" STYLEREF KL一级标题 \n \* MERGEFORMAT ", str(chapter_no))
    caption.add_run("-")
    append_field(caption.add_run(), r" SEQ 表 \* ARABIC \s 1 ", str(table_no))
    caption.add_run(f" {caption_text}")
    blank = insert_paragraph_before(caption, STYLE_BODY)
    normalize_paragraph_to_style(blank, STYLE_BODY)


def make_context_caption(
    heading_1: str | None,
    heading_2: str | None,
    heading_3: str | None,
    figure_no: int = 1,
) -> str:
    """根据附近标题生成缺省图片题注名称。"""
    del figure_no
    for heading in (heading_2, heading_1, heading_3):
        if heading:
            base = heading[:18]
            return base if base.endswith("图") else f"{base}图"
    return "示意图"


def make_context_table_caption(
    heading_1: str | None,
    heading_2: str | None,
    heading_3: str | None,
) -> str:
    """根据附近标题生成缺省表题注名称。"""
    for heading in (heading_2, heading_1, heading_3):
        if heading:
            base = heading[:18]
            return base if base.endswith("表") else f"{base}表"
    return "数据表"


def next_nonblank_block(blocks: list[Paragraph | Table], start: int) -> Paragraph | Table | None:
    """获取后续第一个非空正文块。"""
    for block in blocks[start + 1 :]:
        if isinstance(block, Table):
            return block
        if isinstance(block, Paragraph) and (block.text.strip() or paragraph_has_image(block)):
            return block
    return None


def previous_block(blocks: list[Paragraph | Table], start: int) -> Paragraph | Table | None:
    """获取前一个正文块。"""
    if start <= 0:
        return None
    return blocks[start - 1]


def should_keep_blank(blocks: list[Paragraph | Table], idx: int) -> bool:
    """判断空段落是否是题注规范要求保留的空行。"""
    prev_block = previous_block(blocks, idx)
    next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
    if isinstance(prev_block, Paragraph) and (
        has_figure_caption_field(prev_block) or has_table_caption_field(prev_block)
    ):
        return True
    if isinstance(next_block, Paragraph) and has_table_caption_field(next_block):
        return True
    return False


def is_heading_paragraph(paragraph: Paragraph) -> bool:
    """判断段落是否为标题样式。"""
    name = style_name(paragraph)
    normalized = name.replace(" ", "").lower()
    if name in {
        STYLE_HEADING_1,
        STYLE_HEADING_2,
        STYLE_HEADING_3,
        STYLE_HEADING_4,
        STYLE_HEADING_5,
        STYLE_HEADING_6,
    }:
        return True
    return any(normalized == f"heading{level}" or normalized == f"标题{level}" for level in range(1, 7))


def ensure_blank_before_heading(paragraph: Paragraph) -> None:
    """当前一块不是标题时，确保标题前有一个 KL 正文空段落。"""
    prev = paragraph._element.getprevious()
    if prev is not None and prev.tag == qn("w:p"):
        prev_para = Paragraph(prev, paragraph._parent)
        if is_heading_paragraph(prev_para):
            return
        if not prev_para.text.strip() and not paragraph_has_image(prev_para):
            normalize_paragraph_to_style(prev_para, STYLE_BODY)
            return
    new_p = OxmlElement("w:p")
    paragraph._element.addprevious(new_p)
    blank = Paragraph(new_p, paragraph._parent)
    normalize_paragraph_to_style(blank, STYLE_BODY)


def apply_table_styles(table: Table) -> None:
    """给表格首行和表体应用 KL 表格样式。"""
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                normalize_table_cell_paragraph(
                    paragraph,
                    STYLE_TABLE_HEADER if row_idx == 0 else STYLE_TABLE_BODY,
                )


def _make_source_paragraph_lookup(
    source_blocks: list[dict[str, Any]] | None,
) -> dict[str, deque[dict[str, str]]]:
    """把源文档段落文本映射到源分类，转换阶段优先消费该分类。"""
    lookup: dict[str, deque[dict[str, str]]] = defaultdict(deque)
    if not source_blocks:
        return lookup
    for block in source_blocks:
        if block.get("kind") != "paragraph":
            continue
        block_type = str(block.get("type", ""))
        text = str(block.get("text", "")).strip()
        if not text or block_type in {"blank", "figure"}:
            continue
        item = {
            "type": block_type,
            "reason": str(block.get("reason", "")) or "matched source document classification",
            "confidence": str(block.get("confidence", "")) or "high",
        }
        lookup[text].append(item)
        cleaned = clean_heading_number(text)[0]
        if cleaned and cleaned != text:
            lookup[cleaned].append(item)
    return lookup


def _pop_source_paragraph(
    lookup: dict[str, deque[dict[str, str]]],
    text: str,
    cleaned_text: str,
) -> dict[str, str] | None:
    """按正文顺序取出源文档中同名段落的原始分类。"""
    for key in (text, cleaned_text):
        queue = lookup.get(key)
        if queue:
            return queue.popleft()
    return None


def standardize_appended_body(
    doc: Document,
    start_idx: int,
    source_title: str | None = None,
    source_blocks: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """规范化追加到母版后的正文块。"""
    blocks = iter_block_items(doc)
    report: dict[str, Any] = {
        "styled": [],
        "figure_captions_added": [],
        "table_captions_added": [],
        "captions_standardized": [],
        "tables_styled": 0,
        "removed_duplicate_body_title": False,
        "number_prefixes_removed": [],
    }
    current_chapter_no = 0
    current_heading_1: str | None = None
    current_heading_2: str | None = None
    current_heading_3: str | None = None
    figure_no_by_chapter: dict[int, int] = {}
    table_no_by_chapter: dict[int, int] = {}
    idx = start_idx
    source_paragraph_lookup = _make_source_paragraph_lookup(source_blocks)

    while idx < len(blocks):
        block = blocks[idx]
        if isinstance(block, Table):
            apply_table_styles(block)
            report["tables_styled"] += 1
            prev_block = previous_block(blocks, idx)
            if not (isinstance(prev_block, Paragraph) and is_table_caption_paragraph(prev_block)):
                chapter_no = max(current_chapter_no, 1)
                table_no_by_chapter[chapter_no] = table_no_by_chapter.get(chapter_no, 0) + 1
                table_no = table_no_by_chapter[chapter_no]
                caption_text = make_context_table_caption(
                    current_heading_1,
                    current_heading_2,
                    current_heading_3,
                )
                add_table_caption_before(block, caption_text, chapter_no, table_no)
                report["table_captions_added"].append(
                    {
                        "before_block": idx,
                        "chapter": chapter_no,
                        "table": table_no,
                        "caption": caption_text,
                    }
                )
                blocks = iter_block_items(doc)
                idx += 3
                continue
            idx += 1
            continue

        text = block.text.strip()
        cleaned_text = clean_heading_number(text)[0]
        if (
            source_title
            and cleaned_text == source_title
            and not report["removed_duplicate_body_title"]
        ):
            remove_block(block)
            report["removed_duplicate_body_title"] = True
            blocks = iter_block_items(doc)
            continue

        if paragraph_has_image(block):
            normalize_paragraph_to_style(block, STYLE_IMAGE)
            next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
            if isinstance(next_block, Paragraph) and is_plain_figure_caption_text(next_block.text):
                chapter_no = max(current_chapter_no, 1)
                figure_no_by_chapter[chapter_no] = figure_no_by_chapter.get(chapter_no, 0) + 1
                figure_no = figure_no_by_chapter[chapter_no]
                caption_text = next_block.text.strip()
                set_figure_caption_paragraph(next_block, caption_text, chapter_no, figure_no)
                _pop_source_paragraph(source_paragraph_lookup, caption_text, caption_text)
                report["figure_captions_added"].append(
                    {
                        "after_block": idx,
                        "chapter": chapter_no,
                        "figure": figure_no,
                        "caption": caption_text,
                        "source": "following plain caption text",
                    }
                )
                blocks = iter_block_items(doc)
                idx += 2
                continue
            elif not (isinstance(next_block, Paragraph) and has_figure_caption_field(next_block)):
                chapter_no = max(current_chapter_no, 1)
                figure_no_by_chapter[chapter_no] = figure_no_by_chapter.get(chapter_no, 0) + 1
                figure_no = figure_no_by_chapter[chapter_no]
                caption_text = make_context_caption(
                    current_heading_1,
                    current_heading_2,
                    current_heading_3,
                    figure_no,
                )
                add_figure_caption_after(block, caption_text, chapter_no, figure_no)
                report["figure_captions_added"].append(
                    {
                        "after_block": idx,
                        "chapter": chapter_no,
                        "figure": figure_no,
                        "caption": caption_text,
                    }
                )
                blocks = iter_block_items(doc)
                idx += 3
                continue
            idx += 1
            continue

        if not text:
            if should_keep_blank(blocks, idx):
                normalize_paragraph_to_style(block, STYLE_BODY)
            else:
                remove_block(block)
                blocks = iter_block_items(doc)
                continue
            idx += 1
            continue

        if is_figure_caption_paragraph(block):
            chapter_no = max(current_chapter_no, 1)
            figure_no_by_chapter[chapter_no] = figure_no_by_chapter.get(chapter_no, 0) + 1
            figure_no = figure_no_by_chapter[chapter_no]
            caption_text = extract_caption_title(block, "图") or make_context_caption(
                current_heading_1,
                current_heading_2,
                current_heading_3,
                figure_no,
            )
            was_standard = has_standard_caption_fields(block, "图")
            set_figure_caption_paragraph(block, caption_text, chapter_no, figure_no)
            report["captions_standardized"].append(
                {
                    "block": idx,
                    "kind": "figure",
                    "caption": caption_text,
                    "chapter": chapter_no,
                    "number": figure_no,
                    "was_standard": was_standard,
                }
            )
            idx += 1
            continue

        if is_table_caption_paragraph(block):
            chapter_no = max(current_chapter_no, 1)
            table_no_by_chapter[chapter_no] = table_no_by_chapter.get(chapter_no, 0) + 1
            table_no = table_no_by_chapter[chapter_no]
            caption_text = extract_caption_title(block, "表") or make_context_table_caption(
                current_heading_1,
                current_heading_2,
                current_heading_3,
            )
            was_standard = has_standard_caption_fields(block, "表")
            set_table_caption_paragraph(block, caption_text, chapter_no, table_no)
            report["captions_standardized"].append(
                {
                    "block": idx,
                    "kind": "table",
                    "caption": caption_text,
                    "chapter": chapter_no,
                    "number": table_no,
                    "was_standard": was_standard,
                }
            )
            idx += 1
            continue

        source_paragraph = _pop_source_paragraph(source_paragraph_lookup, text, cleaned_text)
        if source_paragraph:
            source_type = source_paragraph["type"]
            if source_type.startswith("heading_"):
                kind = source_type
                clean_text = clean_heading_number(text)[0]
            elif source_type in {"figure_caption", "table_caption"}:
                normalize_paragraph_to_style(block, STYLE_CAPTION)
                idx += 1
                continue
            else:
                kind = "body"
                clean_text = text
            reason = source_paragraph["reason"]
            confidence = source_paragraph["confidence"]
        else:
            kind, clean_text, reason, confidence = classify_heading(text, block)
        following = next_nonblank_block(blocks, idx)
        if (
            kind == "heading_2"
            and reason == "short standalone title line"
            and isinstance(following, Paragraph)
            and paragraph_has_image(following)
        ):
            kind = "body"
            clean_text = text
            reason = "short image context line preserved as body"
            confidence = "medium"

        if kind.startswith("heading_"):
            clean_text = normalize_heading_text(clean_text)
            clear_numbering(block)
            if clean_text != block.text:
                original_text = block.text
                set_paragraph_text(block, clean_text)
                report["number_prefixes_removed"].append(
                    {"block": idx, "from": original_text, "to": clean_text}
                )

        if kind == "heading_1":
            ensure_blank_before_heading(block)
            current_chapter_no += 1
            current_heading_1 = clean_text
            current_heading_2 = None
            current_heading_3 = None
            normalize_paragraph_to_style(block, STYLE_HEADING_1)
        elif kind == "heading_2":
            ensure_blank_before_heading(block)
            current_heading_2 = clean_text
            current_heading_3 = None
            normalize_paragraph_to_style(block, STYLE_HEADING_2)
        elif kind == "heading_3":
            ensure_blank_before_heading(block)
            current_heading_3 = clean_text
            normalize_paragraph_to_style(block, STYLE_HEADING_3)
        elif kind == "heading_4":
            ensure_blank_before_heading(block)
            normalize_paragraph_to_style(block, STYLE_HEADING_4)
        elif kind == "heading_5":
            ensure_blank_before_heading(block)
            normalize_paragraph_to_style(block, STYLE_HEADING_5)
        elif kind == "heading_6":
            ensure_blank_before_heading(block)
            normalize_paragraph_to_style(block, STYLE_HEADING_6)
        else:
            normalize_paragraph_to_style(block, STYLE_BODY)

        report["styled"].append(
            {
                "block": idx,
                "type": kind,
                "text": block.text.strip(),
                "reason": reason,
                "confidence": confidence,
            }
        )
        idx += 1

    return report


def build_standard_docx(
    input_docx: Path | str,
    output_docx: Path | str,
    template_docx: Path | str | None = None,
    *,
    skip_source_front_matter: bool = False,
) -> dict[str, Any]:
    """基于 KL 母版把正文文档转换为标准文档。"""
    input_path = Path(input_docx)
    output_path = Path(output_docx)
    template_path = Path(template_docx) if template_docx is not None else DEFAULT_MASTER_TEMPLATE
    if not template_path.exists():
        raise FileNotFoundError(f"KL母版模板不存在: {template_path}")
    source_report = classify_docx_body(input_path)
    detected_source_front_matter = detect_source_front_matter(source_report)
    source_front_matter = (
        detected_source_front_matter
        if skip_source_front_matter
        else {
            "detected": False,
            "body_start_block": 0,
            "body_start_reason": "source is treated as body content from the first block",
            "skipped_blocks_count": 0,
            "skipped_blocks": [],
            "auto_detected": detected_source_front_matter,
        }
    )
    source_title = infer_source_title(input_path, source_report, source_front_matter)
    source_file_number = infer_source_file_number(source_report)
    source_body_start = int(source_front_matter["body_start_block"])
    source_blocks = source_report["blocks"][source_body_start:]
    has_figures = document_has_figures(input_path)
    has_tables = document_has_tables(input_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        front_docx = tmp_dir / "front.docx"
        shutil.copy2(template_path, front_docx)

        front_doc = Document(str(front_docx))
        replace_cover_title(front_doc, source_title)
        replace_cover_file_number(front_doc, source_file_number)
        remove_template_front_matter_samples(front_doc)
        set_default_headers(front_doc, source_title)
        front_matter = {"template_front_matter_preserved": True}
        front_count = remove_template_sample_body(front_doc)
        front_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        front_doc.save(str(front_docx))

        source_body_docx = tmp_dir / "source_body.docx"
        write_source_body_only_docx(input_path, source_body_docx, source_body_start)

        master = Document(str(front_docx))
        composer = Composer(master)
        composer.append(Document(str(source_body_docx)))
        composed_docx = tmp_dir / "composed.docx"
        composer.save(str(composed_docx))

        doc = Document(str(composed_docx))
        set_default_headers(doc, source_title)
        report = standardize_appended_body(doc, front_count + 1, source_title, source_blocks)
        doc.save(str(output_path))

    report.update(
        {
            "input": str(input_path),
            "output": str(output_path),
            "template": str(template_path),
            "source_title": source_title,
            "source_file_number": source_file_number,
            "input_has_figures": has_figures,
            "input_has_tables": has_tables,
            "front_matter": front_matter,
            "front_blocks_preserved": front_count,
            "source_front_matter": source_front_matter,
            "skip_source_front_matter": skip_source_front_matter,
        }
    )
    return report


def _xml_qn(tag: str) -> str:
    prefix, local = tag.split(":", 1)
    return f"{{{NS[prefix]}}}{local}"


def read_docx_xml(docx: Path | str, name: str) -> str:
    """读取 DOCX 包中的 XML 文件文本。"""
    with zipfile.ZipFile(Path(docx)) as zf:
        try:
            return zf.read(name).decode("utf-8", errors="replace")
        except KeyError:
            return ""


def load_style_names(docx: Path | str) -> dict[str, str]:
    """读取 styles.xml 中 styleId 到样式名称的映射。"""
    styles_xml = read_docx_xml(docx, "word/styles.xml")
    if not styles_xml:
        return {}
    root = ET.fromstring(styles_xml)
    result: dict[str, str] = {}
    for style in root.findall("w:style", NS):
        style_id = style.attrib.get(_xml_qn("w:styleId"))
        name_node = style.find("w:name", NS)
        if not style_id or name_node is None:
            continue
        style_value = name_node.attrib.get(_xml_qn("w:val"))
        if style_value:
            result[style_id] = style_value
    return result


def _paragraph_text_xml(p: ET.Element) -> str:
    return "".join(t.text or "" for t in p.findall(".//w:t", NS)).strip()


def _paragraph_style_id_xml(p: ET.Element) -> str | None:
    pstyle = p.find("./w:pPr/w:pStyle", NS)
    if pstyle is None:
        return None
    return pstyle.attrib.get(_xml_qn("w:val"))


def _paragraph_field_codes_xml(p: ET.Element) -> list[str]:
    codes = []
    for node in p.findall(".//w:instrText", NS):
        if node.text and node.text.strip():
            codes.append(node.text.strip())
    return codes


def _has_figure_xml(p: ET.Element) -> bool:
    return p.find(".//w:drawing", NS) is not None or p.find(".//w:pict", NS) is not None


def _classify_paragraph_xml(p: ET.Element, style_names: dict[str, str]) -> tuple[str, str, str]:
    text = _paragraph_text_xml(p)
    style_id = _paragraph_style_id_xml(p)
    style_value = style_names.get(style_id or "", style_id or "")
    fields = "\n".join(_paragraph_field_codes_xml(p))

    if "SEQ 图" in fields or CAPTION_TEXT_PATTERNS["figure_caption"].match(text):
        return "figure_caption", "caption field/text matched figure", "high"
    if "SEQ 表" in fields or CAPTION_TEXT_PATTERNS["table_caption"].match(text):
        return "table_caption", "caption field/text matched table", "high"
    if _has_figure_xml(p):
        return "figure", "paragraph contains drawing or pict", "high"
    normalized_style = (style_value or "").replace(" ", "").lower()
    if style_value in STYLE_TO_TYPE:
        kind, confidence = STYLE_TO_TYPE[style_value]
        return kind, f"matched style {style_value}", confidence
    for level in range(1, 7):
        if normalized_style in {f"heading{level}", f"标题{level}", str(level)}:
            return f"heading_{level}", f"matched style {style_value}", "high"
    if normalized_style == "listparagraph":
        return "body", "list paragraph is body, not standalone heading", "high"
    kind, _cleaned, reason, confidence = classify_heading(text)
    return (
        kind if text else "blank",
        reason if text else "empty paragraph",
        confidence if text else "high",
    )


def _table_dimensions_xml(tbl: ET.Element) -> tuple[int, int]:
    rows = tbl.findall("./w:tr", NS)
    max_cols = 0
    for row in rows:
        max_cols = max(max_cols, len(row.findall("./w:tc", NS)))
    return len(rows), max_cols


def _find_table_caption(blocks: list[BlockReport], table_index: int) -> dict[str, Any] | None:
    before = blocks[table_index - 1] if table_index > 0 else None
    after = blocks[table_index + 1] if table_index + 1 < len(blocks) else None
    if before and before.type == "table_caption":
        return {"position": "above", "block_index": before.index, "text": before.text}
    if after and after.type == "table_caption":
        return {"position": "below", "block_index": after.index, "text": after.text}
    return None


def _find_figure_caption(blocks: list[BlockReport], figure_index: int) -> dict[str, Any] | None:
    before = blocks[figure_index - 1] if figure_index > 0 else None
    after = blocks[figure_index + 1] if figure_index + 1 < len(blocks) else None
    if after and after.type == "figure_caption":
        return {"position": "below", "block_index": after.index, "text": after.text}
    if before and before.type == "figure_caption":
        return {"position": "above", "block_index": before.index, "text": before.text}
    return None


def classify_docx_body(docx: Path | str) -> dict[str, Any]:
    """分类 DOCX 正文结构，返回标题、正文、图、表、题注报告。"""
    docx_path = Path(docx)
    document_xml = read_docx_xml(docx_path, "word/document.xml")
    if not document_xml:
        raise ValueError("word/document.xml not found")

    style_names = load_style_names(docx_path)
    root = ET.fromstring(document_xml)
    body = root.find("w:body", NS)
    if body is None:
        raise ValueError("document body not found")

    blocks: list[BlockReport] = []
    for child in list(body):
        if child.tag == _xml_qn("w:p"):
            style_id = _paragraph_style_id_xml(child)
            style_value = style_names.get(style_id or "", style_id)
            field_codes = _paragraph_field_codes_xml(child)
            block_type, reason, confidence = _classify_paragraph_xml(child, style_names)
            blocks.append(
                BlockReport(
                    index=len(blocks),
                    kind="paragraph",
                    type=block_type,
                    text=_paragraph_text_xml(child),
                    style_id=style_id,
                    style_name=style_value,
                    reason=reason,
                    confidence=confidence,
                    has_figure=_has_figure_xml(child),
                    field_codes=field_codes or None,
                )
            )
        elif child.tag == _xml_qn("w:tbl"):
            rows, columns = _table_dimensions_xml(child)
            blocks.append(
                BlockReport(
                    index=len(blocks),
                    kind="table",
                    type="table",
                    rows=rows,
                    columns=columns,
                    reason="block is w:tbl",
                    confidence="high",
                )
            )

    for idx, block in enumerate(blocks):
        if block.type == "table":
            block.caption = _find_table_caption(blocks, idx)
        elif block.type == "figure":
            block.caption = _find_figure_caption(blocks, idx)

    summary: dict[str, int] = {}
    for block in blocks:
        summary[block.type] = summary.get(block.type, 0) + 1

    return {
        "source": str(docx_path),
        "summary": summary,
        "blocks": [asdict(block) for block in blocks],
        "headings": [asdict(block) for block in blocks if block.type.startswith("heading_")],
        "figures": [asdict(block) for block in blocks if block.type == "figure"],
        "tables": [asdict(block) for block in blocks if block.type == "table"],
        "captions": [
            asdict(block)
            for block in blocks
            if block.type in {"figure_caption", "table_caption"}
        ],
    }


def write_classification_report(
    docx: Path | str,
    output_json: Path | str,
    *,
    pretty: bool = True,
) -> Path:
    """把 DOCX 分类结果写入 JSON 文件。"""
    output = Path(output_json)
    output.parent.mkdir(parents=True, exist_ok=True)
    report = classify_docx_body(docx)
    output.write_text(
        json.dumps(report, ensure_ascii=False, indent=2 if pretty else None) + "\n",
        encoding="utf-8",
    )
    return output


def style_names_in_docx(docx: Path | str) -> set[str]:
    """读取文档中定义的样式名称集合。"""
    styles_xml = read_docx_xml(docx, "word/styles.xml")
    if not styles_xml:
        return set()
    root = ET.fromstring(styles_xml)
    names: set[str] = set()
    for style in root.findall("w:style", NS):
        name_node = style.find("w:name", NS)
        if name_node is not None:
            value = name_node.attrib.get(_xml_qn("w:val"))
            if value:
                names.add(value)
    return names


def field_instr_text(docx: Path | str) -> str:
    """读取 document.xml 中所有域代码文本。"""
    document_xml = read_docx_xml(docx, "word/document.xml")
    if not document_xml:
        return ""
    root = ET.fromstring(document_xml)
    parts = [node.text for node in root.findall(".//w:instrText", NS) if node.text]
    return "\n".join(parts)


def check_word_standard(docx: Path | str) -> dict[str, Any]:
    """检查 DOCX 是否包含 KL 标准核心样式和题注/清单字段信号。"""
    docx_path = Path(docx)
    found_styles = style_names_in_docx(docx_path)
    missing_styles = [name for name in REQUIRED_STYLES if name not in found_styles]
    fields = field_instr_text(docx_path)
    missing_fields: list[str] = []
    if ("SEQ 图" in fields or "SEQ 表" in fields) and "STYLEREF KL一级标题" not in fields:
        missing_fields.append("STYLEREF KL一级标题")
    missing_toc = [
        snippet
        for caption, snippet in {
            "SEQ 图": 'TOC \\h \\z \\c "图"',
            "SEQ 表": 'TOC \\h \\z \\c "表"',
        }.items()
        if caption in fields and snippet not in fields
    ]
    return {
        "source": str(docx_path),
        "found_styles": sorted(found_styles),
        "required_styles_ok": not missing_styles,
        "missing_styles": missing_styles,
        "caption_fields_ok": not missing_fields,
        "missing_field_snippets": missing_fields,
        "figure_table_lists_ok": not missing_toc,
        "missing_toc_snippets": missing_toc,
        "has_wrong_table_toc": 'TOC \\h \\z \\c "表格"' in fields,
        "has_table_seq": "SEQ 表" in fields,
        "has_figure_seq": "SEQ 图" in fields,
    }


__all__ = [
    "STYLE_MAIN_TITLE",
    "STYLE_HEADING_1",
    "STYLE_HEADING_2",
    "STYLE_HEADING_3",
    "STYLE_HEADING_4",
    "STYLE_HEADING_5",
    "STYLE_HEADING_6",
    "STYLE_BODY",
    "STYLE_CAPTION",
    "STYLE_TABLE_HEADER",
    "STYLE_TABLE_BODY",
    "STYLE_IMAGE",
    "PACKAGE_ASSETS_DIR",
    "DEFAULT_MASTER_TEMPLATE",
    "REQUIRED_STYLES",
    "BlockReport",
    "iter_block_items",
    "clean_heading_number",
    "classify_heading",
    "normalize_paragraph_to_style",
    "normalize_table_cell_paragraph",
    "apply_table_styles",
    "add_figure_caption_after",
    "detect_source_front_matter",
    "classify_docx_body",
    "write_classification_report",
    "check_word_standard",
    "build_standard_docx",
]
